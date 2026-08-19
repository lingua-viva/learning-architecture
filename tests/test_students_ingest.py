"""T9 — Students-from-file ingest (SPEC_T9_INGEST_UI_2026-08-04).

The T3 extraction seam is exercised with the T0 fixture extraction (the
frozen contract output shape); one test pins the honest failure while the
real stub still raises NotImplementedError.
"""
from __future__ import annotations

import asyncio
import json
import time
from io import BytesIO
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

import src.web as web
from src.lingua_viva.docpipe import vault
from src.lingua_viva.docpipe.contracts import ExtractionRecord

REPO = Path(__file__).resolve().parents[1]
FIXTURES = REPO / "tests" / "fixtures" / "docpipe"
client = TestClient(web.app)


@pytest.fixture()
def isolated_state(tmp_path, monkeypatch):
    monkeypatch.setenv("LV_STATE_HOME", str(tmp_path))
    monkeypatch.setenv("LV_STUDENT_DB_PATH", str(tmp_path / "student_lenses.db"))
    monkeypatch.setattr(web, "_INGEST_JOBS", {})
    return tmp_path


def _fixture_extraction(source_id: str) -> dict:
    data = json.loads((FIXTURES / "expected_extraction_lesson_plan_marco_nora.json").read_text(encoding="utf-8"))
    data["source_id"] = source_id
    # The golden fixture carries the real pipeline's evidence for this markdown
    # file (bigram_fallback). Ingest tests simulate a TRUSTED structured source
    # by default; tests about the untrusted path override this per student
    # (STEP 3, SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19).
    for student in data["structure"]["students_detected"]:
        student["evidence"] = "student_column"
    return data


def _roster_extraction(source_id: str, names: list[str]) -> dict:
    text = "Class List\n" + "\n".join(names) + "\n\n"
    spans = [
        {
            "span_id": "SPN-0001",
            "char_start": 0,
            "char_end": len(text.strip()),
            "text": text.strip(),
        }
    ]
    return {
        "schema_version": "docpipe.extraction.v1",
        "source_id": source_id,
        "source_sha256": "test",
        "extracted_at": "2026-08-06T00:00:00Z",
        "extractor": {"name": "fixture", "version": "1", "model": None},
        "mime": "text/markdown",
        "language": "en",
        "normalized_text": text,
        "spans": spans,
        "structure": {
            "title": "Class List",
            "document_type": "roster",
            "sections": [],
            "students_detected": [
                {
                    "student_id": f"student-{name.lower().replace(' ', '-')}",
                    "display_name": name,
                    "confidence": 0.99,
                    "evidence": "student_column",
                    "span_ids": ["SPN-0001"],
                }
                for name in names
            ],
        },
        "warnings": [],
    }


@pytest.fixture()
def fixture_extractor(monkeypatch):
    """Stand-in for T3 with the frozen contract output shape."""
    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)


def _upload(name: str = "lesson_plan_marco_nora.md") -> dict:
    content = (FIXTURES / name).read_bytes()
    response = client.post(
        "/api/students/ingest",
        files={"file": (name, content, "text/markdown")},
    )
    assert response.status_code == 200, response.text
    return response.json()


def _upload_bytes(name: str, content: bytes, mime: str) -> dict:
    response = client.post(
        "/api/students/ingest",
        files={"file": (name, content, mime)},
    )
    assert response.status_code == 200, response.text
    return response.json()


def _wait_for_job(job_id: str, timeout: float = 10.0) -> dict:
    deadline = time.time() + timeout
    while time.time() < deadline:
        response = client.get(f"/api/students/ingest/{job_id}")
        assert response.status_code == 200, response.text
        job = response.json()
        if job["status"] in ("done", "failed", "preview", "cancelled"):
            return job
        time.sleep(0.05)
    raise AssertionError("ingest job did not finish in time")


def _run_to_done(job_id: str, timeout: float = 10.0) -> dict:
    """Phase 0A (SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19 §2A): every import
    with detections stops at a preview; creation needs the explicit approve.
    This helper drives that two-step flow to completion for tests whose
    subject is what happens AFTER creation. Jobs that end at done/failed
    without a preview (zero detections, extraction failure) pass through."""
    job = _wait_for_job(job_id, timeout)
    if job["status"] != "preview":
        return job
    response = client.post("/api/students/ingest/approve", json={"job_id": job_id})
    assert response.status_code == 200, response.text
    return _wait_for_job(job_id, timeout)


def _confirm_pending(job: dict) -> dict:
    """STEP 3 (SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19): names detected by
    the prose fallback (bigram_fallback evidence) require per-name
    confirmation on small imports. Drive the confirm step for everything
    pending and return the refreshed job."""
    names = [item["display_name"] for item in job.get("needs_confirmation", [])]
    if names:
        response = client.post("/api/students/ingest/confirm", json={
            "job_id": job["job_id"], "display_names": names,
        })
        assert response.status_code == 200, response.text
    return client.get(f"/api/students/ingest/{job['job_id']}").json()


def test_full_chain_creates_grounded_lenses(isolated_state, fixture_extractor):
    started = _upload()
    job = _run_to_done(started["job_id"])
    assert job["status"] == "done", job
    assert job["students_found"] == 2
    names = sorted(s["display_name"] for s in job["students_created"])
    assert names == ["Marco Bianchi", "Nora Rossi"]

    # Vault artifacts validate against the frozen schemas.
    from src.lingua_viva.docpipe.validate import validate_file

    root = vault.vault_root()
    assert not validate_file(root / "manifest.json")
    for created in job["students_created"]:
        lens_path = root / "lenses" / created["student_id"] / "lens.json"
        assert not validate_file(lens_path)
        lens = json.loads(lens_path.read_text(encoding="utf-8"))
        populated = [f for f in lens["profile"].values() if f["evidence"]]
        assert populated, "created lens has no grounded fields"
        for field in lens["profile"].values():
            if field["value"] in (None, "", [], {}):
                assert field["evidence"] == []
            else:
                assert field["evidence"], "populated field without evidence"

    # Bridge: the students appear in the existing roster store.
    def roster(store):
        return sorted(str(l.get("display_name")) for l in store.list_lenses())

    names_in_store = web._with_student_store(roster)
    assert "Marco Bianchi" in names_in_store and "Nora Rossi" in names_in_store

    # Operator ruling §8.1: every locally saved lens is queued for Drive
    # write-back (the T6 queue holds until push_file can drain).
    from src.lingua_viva.docpipe import sync as docpipe_sync

    for created in job["students_created"]:
        assert docpipe_sync.sync_status(created["student_id"])["status"] in ("pending", "failed")


def test_reimport_merges_never_forks(isolated_state, fixture_extractor):
    first = _run_to_done(_upload()["job_id"])
    assert first["status"] == "done"
    root = vault.vault_root()
    marco_id = next(s["student_id"] for s in first["students_created"] if s["display_name"] == "Marco Bianchi")
    before = json.loads((root / "lenses" / marco_id / "lens.json").read_text(encoding="utf-8"))
    evidence_before = sum(len(f["evidence"]) for f in before["profile"].values())

    second = _run_to_done(_upload()["job_id"])
    assert second["status"] == "done"
    after = json.loads((root / "lenses" / marco_id / "lens.json").read_text(encoding="utf-8"))
    evidence_after = sum(len(f["evidence"]) for f in after["profile"].values())
    # Different source_id per import → new evidence keys are allowed, but the
    # student must NOT fork into a second lens.
    lens_dirs = [p.name for p in (root / "lenses").iterdir() if p.is_dir()]
    assert lens_dirs.count(marco_id) == 1
    assert len([d for d in lens_dirs if "marco" in d.lower()]) == 1
    assert evidence_after >= evidence_before


def test_low_confidence_student_needs_confirmation(isolated_state, monkeypatch):
    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        for student in data["structure"]["students_detected"]:
            student["evidence"] = "bigram_fallback"
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _run_to_done(_upload()["job_id"])
    assert job["status"] == "done"
    assert job["students_created"] == []
    assert len(job["needs_confirmation"]) == 2

    # Teacher confirms one — only then is the lens created.
    response = client.post("/api/students/ingest/confirm", json={
        "job_id": job["job_id"], "display_name": "Nora Rossi",
    })
    assert response.status_code == 200, response.text
    created = response.json()["student"]
    assert created["display_name"] == "Nora Rossi"
    lens_path = vault.vault_root() / "lenses" / created["student_id"] / "lens.json"
    assert lens_path.exists()
    refreshed = client.get(f"/api/students/ingest/{job['job_id']}").json()
    assert all(item["display_name"] != "Nora Rossi" for item in refreshed["needs_confirmation"])


def test_small_import_confirm_multiple_names_at_once(isolated_state, monkeypatch):
    """The confirm endpoint's display_names batch form still serves the
    small-import path (≤2 students detected without a student-name column)."""
    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        for student in data["structure"]["students_detected"]:
            student["evidence"] = "bigram_fallback"
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _run_to_done(_upload()["job_id"])
    assert len(job["needs_confirmation"]) == 2
    names = [item["display_name"] for item in job["needs_confirmation"]]

    response = client.post("/api/students/ingest/confirm", json={
        "job_id": job["job_id"], "display_names": names,
    })
    assert response.status_code == 200, response.text
    created = response.json()["students"]
    assert [student["display_name"] for student in created] == names
    refreshed = client.get(f"/api/students/ingest/{job['job_id']}").json()
    assert refreshed["needs_confirmation"] == []


def test_bulk_roster_auto_creates_every_student(isolated_state, monkeypatch):
    """SPEC_LV_DRIVE_OOTB 2026-08-18 (G3): roster imports create ALL students
    with zero confirm clicks — review happens via one-click undo, not per-name
    approval. This inverted the pre-spec behavior (roster → needs_confirmation)."""
    names = ["Marco Bianchi", "Nora Rossi", "Luca Verdi", "Sara Conti", "Maya Singh"]

    async def extract(source, content, *, model_client=None):
        data = _roster_extraction(source.source_id, names)
        data["source_sha256"] = source.data["sha256"]
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _run_to_done(_upload()["job_id"])

    assert job["status"] == "done"
    assert job["students_found"] == 5
    assert [s["display_name"] for s in job["students_created"]] == names
    assert job["needs_confirmation"] == []

    def roster(store):
        return sorted(lens["display_name"] for lens in store.list_lenses())

    assert web._with_student_store(roster) == sorted(names)
    for created in job["students_created"]:
        lens_path = vault.vault_root() / "lenses" / created["student_id"] / "lens.json"
        assert lens_path.exists()


def test_bulk_roster_low_confidence_names_created_and_flagged(isolated_state, monkeypatch):
    """Untrusted roster names (no student-name column behind them) are still
    created (no confirm clicks) but flagged in warnings so the teacher knows
    which names to double-check."""
    names = ["Marco Bianchi", "Nora Rossi", "Luca Verdi", "Sara Conti", "Maya Singh"]

    async def extract(source, content, *, model_client=None):
        data = _roster_extraction(source.source_id, names)
        data["source_sha256"] = source.data["sha256"]
        for student in data["structure"]["students_detected"]:
            if student["display_name"] in ("Nora Rossi", "Maya Singh"):
                student["evidence"] = "bigram_fallback"
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _run_to_done(_upload()["job_id"])

    assert job["status"] == "done"
    assert [s["display_name"] for s in job["students_created"]] == names
    assert job["needs_confirmation"] == []
    flagged = [w for w in job["warnings"] if "Check these names" in w]
    assert len(flagged) == 1
    assert "Nora Rossi" in flagged[0] and "Maya Singh" in flagged[0]
    assert "Marco Bianchi" not in flagged[0]


def test_gate_rides_on_evidence_class_not_confidence_number(isolated_state, monkeypatch):
    """STEP 3 class lock (SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19): the old
    numeric gate was dead code — confidence is a flat constant, so
    `confidence >= threshold` could never fail. Trust now rides on the
    corpus-measured evidence class: a fallback guess with a HIGH number still
    needs confirmation, and a detection with no evidence class at all is
    untrusted. No number can buy trust, and the dead constant stays deleted."""
    assert not hasattr(web, "INGEST_CONFIDENCE_THRESHOLD")

    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        students = data["structure"]["students_detected"]
        students[0]["evidence"] = "bigram_fallback"
        students[0]["confidence"] = 0.99  # a high number buys nothing
        del students[1]["evidence"]  # unknown origin — untrusted
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _wait_for_job(_upload()["job_id"])
    assert job["status"] == "preview"
    assert all(s["low_confidence"] is True for s in job["preview_students"])

    response = client.post("/api/students/ingest/approve", json={"job_id": job["job_id"]})
    assert response.status_code == 200, response.text
    job = _wait_for_job(job["job_id"])
    assert job["status"] == "done"
    assert job["students_created"] == []
    assert len(job["needs_confirmation"]) == 2


def test_bulk_undo_archives_only_students_from_that_import(isolated_state, monkeypatch):
    existing = client.post("/api/students", json={"display_name": "Existing Student", "grade_level": "G3"})
    assert existing.status_code == 200, existing.text
    names = ["Marco Bianchi", "Nora Rossi"]

    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        data["structure"]["students_detected"] = [
            {
                "student_id": f"student-{name.lower().replace(' ', '-')}",
                "display_name": name,
                "confidence": 0.99,
                "evidence": "student_column",
                "span_ids": ["SPN-0001"],
            }
            for name in names
        ]
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _run_to_done(_upload()["job_id"])
    assert sorted(s["display_name"] for s in job["students_created"]) == names

    response = client.delete(f"/api/students/ingest/{job['job_id']}")
    assert response.status_code == 200, response.text
    assert sorted(s["display_name"] for s in response.json()["students_archived"]) == names

    roster = client.get("/api/students").json()["students"]
    assert [student["display_name"] for student in roster] == ["Existing Student"]


def test_bulk_undo_does_not_touch_students_from_another_import(isolated_state, monkeypatch):
    batches = [
        ["Marco Bianchi", "Nora Rossi"],
        ["Luca Verdi", "Sara Conti"],
    ]

    async def extract(source, content, *, model_client=None):
        names = batches.pop(0)
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        data["structure"]["students_detected"] = [
            {
                "student_id": f"student-{name.lower().replace(' ', '-')}",
                "display_name": name,
                "confidence": 0.99,
                "evidence": "student_column",
                "span_ids": ["SPN-0001"],
            }
            for name in names
        ]
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    first = _run_to_done(_upload()["job_id"])
    second = _run_to_done(_upload()["job_id"])
    assert sorted(s["display_name"] for s in first["students_created"]) == [
        "Marco Bianchi", "Nora Rossi",
    ]
    assert sorted(s["display_name"] for s in second["students_created"]) == [
        "Luca Verdi", "Sara Conti",
    ]

    response = client.delete(f"/api/students/ingest/{first['job_id']}")
    assert response.status_code == 200, response.text

    roster = sorted(student["display_name"] for student in client.get("/api/students").json()["students"])
    assert roster == ["Luca Verdi", "Sara Conti"]


def test_ingest_job_status_survives_memory_reset(isolated_state, fixture_extractor):
    job = _run_to_done(_upload()["job_id"])
    web._INGEST_JOBS = {}

    response = client.get(f"/api/students/ingest/{job['job_id']}")

    assert response.status_code == 200, response.text
    assert response.json()["job_id"] == job["job_id"]


def test_real_extraction_end_to_end_no_mock(isolated_state, monkeypatch):
    """T3 landed: the REAL extractor (deterministic core, no model) drives the
    chain — upload the fixture document, get real grounded lenses. This test
    replaced the honest-stub-failure test the moment T3 shipped."""
    # Force the deterministic path: no model client in the ingest job.
    import src.lingua_viva.docpipe.model as docpipe_model

    monkeypatch.setattr(
        docpipe_model, "LocalModelClient",
        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("no model in test")),
    )
    job = _run_to_done(_upload()["job_id"])
    assert job["status"] == "done", job
    # Prose markdown → bigram_fallback evidence → the real flow asks for
    # per-name confirmation before creating anything (STEP 3): the corpus
    # measured the fallback at precision 0.14, so it never auto-creates.
    assert job["students_created"] == []
    pending = sorted(item["display_name"] for item in job["needs_confirmation"])
    assert pending == ["Marco Bianchi", "Nora Rossi"]
    job = _confirm_pending(job)
    names = sorted(s["display_name"] for s in job["students_created"])
    assert names == ["Marco Bianchi", "Nora Rossi"]
    root = vault.vault_root()
    extraction = json.loads(
        (root / "extracted" / f"{job['source_id']}.json").read_text(encoding="utf-8")
    )
    assert extraction["structure"]["document_type"] == "lesson_plan"
    for created in job["students_created"]:
        lens = json.loads(
            (root / "lenses" / created["student_id"] / "lens.json").read_text(encoding="utf-8")
        )
        assert any(f["evidence"] for f in lens["profile"].values())


def test_real_xlsx_roster_upload_auto_creates_students(isolated_state, monkeypatch):
    """Acceptance: real Excel class list uploads through /api/students/ingest.

    Five detected students is roster-shaped: every student is created
    automatically (SPEC_LV_DRIVE_OOTB G3) — undo is the review mechanism.
    """
    from openpyxl import Workbook

    import src.lingua_viva.docpipe.model as docpipe_model

    monkeypatch.setattr(
        docpipe_model, "LocalModelClient",
        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("no model in test")),
    )
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["First", "Last"])
    names = [
        ("Marco", "Bianchi"),
        ("Nora", "Rossi"),
        ("Luca", "Verdi"),
        ("Sara", "Conti"),
        ("Maya", "Singh"),
    ]
    for first, last in names:
        sheet.append([first, last])
    buffer = BytesIO()
    workbook.save(buffer)

    started = _upload_bytes(
        "class-list.xlsx",
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    job = _run_to_done(started["job_id"])

    assert job["status"] == "done", job
    assert job["students_found"] == 5
    assert [s["display_name"] for s in job["students_created"]] == [
        "Marco Bianchi",
        "Nora Rossi",
        "Luca Verdi",
        "Sara Conti",
        "Maya Singh",
    ]
    assert job["needs_confirmation"] == []


def test_real_docx_roster_upload_auto_creates_students(isolated_state, monkeypatch):
    """Acceptance: real Word table uploads through /api/students/ingest."""
    import docx
    import src.lingua_viva.docpipe.model as docpipe_model

    monkeypatch.setattr(
        docpipe_model, "LocalModelClient",
        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("no model in test")),
    )
    document = docx.Document()
    document.add_paragraph("Class List")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "First"
    table.rows[0].cells[1].text = "Last"
    for first, last in [
        ("Marco", "Bianchi"),
        ("Nora", "Rossi"),
        ("Luca", "Verdi"),
        ("Sara", "Conti"),
        ("Maya", "Singh"),
    ]:
        cells = table.add_row().cells
        cells[0].text = first
        cells[1].text = last
    buffer = BytesIO()
    document.save(buffer)

    started = _upload_bytes(
        "class-list.docx",
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    job = _run_to_done(started["job_id"])

    assert job["status"] == "done", job
    assert job["students_found"] == 5
    assert len(job["students_created"]) == 5
    assert job["needs_confirmation"] == []


def test_messy_roster_fills_only_evidenced_fields(isolated_state, monkeypatch):
    """SPEC_LV_DRIVE_OOTB 2026-08-18 (G4): tolerant fetch-and-fill, locked.

    A messy roster — mixed separators, prose notes, missing fields,
    inconsistent formatting — through the REAL deterministic extractor:
    every student gets a lens; profile fields populate ONLY where the
    document carries evidence; nothing is invented for the students whose
    rows have no usable notes.
    """
    import src.lingua_viva.docpipe.model as docpipe_model

    monkeypatch.setattr(
        docpipe_model, "LocalModelClient",
        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("no model in test")),
    )
    messy = (
        "Class 5B student list (updated!!)\n\n"
        "Marco Bianchi ; benefits from a checklist before writing ; likes football\n\n"
        "Nora Rossi, strong inference skills, can extend ideas with a relevant quotation\n\n"
        "Luca Verdi\n\n"
        "Sara Conti -- new arrival, no notes yet\n\n"
        "(remember to update this file)\n"
    )
    job = _run_to_done(
        _upload_bytes("messy-roster.md", messy.encode("utf-8"), "text/markdown")["job_id"]
    )

    assert job["status"] == "done", job
    created = {s["display_name"]: s for s in job["students_created"]}
    assert sorted(created) == ["Luca Verdi", "Marco Bianchi", "Nora Rossi", "Sara Conti"]
    assert job["needs_confirmation"] == []

    # Fields populate where the row carries evidence…
    assert "executive_functioning" in created["Marco Bianchi"]["fields_populated"]
    assert "strategies_trialed" in created["Marco Bianchi"]["fields_populated"]
    assert "academic_strengths" in created["Nora Rossi"]["fields_populated"]
    # …and stay empty where it doesn't: bare name / prose with no signal.
    assert created["Luca Verdi"]["fields_populated"] == []
    assert created["Sara Conti"]["fields_populated"] == []

    # Zero invented values: every populated field carries document evidence
    # with a span_id, every value is verbatim from the roster, every empty
    # field has no evidence.
    root = vault.vault_root()
    for name, entry in created.items():
        lens = json.loads(
            (root / "lenses" / entry["student_id"] / "lens.json").read_text(encoding="utf-8")
        )
        for field_id, field in lens["profile"].items():
            value = field.get("value")
            if value in (None, "", [], {}):
                assert field.get("evidence", []) == [], (name, field_id)
                continue
            assert field["evidence"], (name, field_id)
            assert all(e.get("span_id") for e in field["evidence"]), (name, field_id)
            values = value if isinstance(value, list) else [value]
            for item in values:
                assert isinstance(item, str) and item in " ".join(messy.split()), (name, field_id, item)


def test_support_document_updates_existing_lenses_no_duplicates(isolated_state, monkeypatch):
    """Operator request 2026-08-18: a support-team document about existing
    students (special needs report, specialist notes) imported later must
    MERGE new evidence into their existing lenses — never create duplicate
    students. Match is by exact name through the same import flow."""
    import src.lingua_viva.docpipe.model as docpipe_model

    monkeypatch.setattr(
        docpipe_model, "LocalModelClient",
        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("no model in test")),
    )
    # Day 1: bare roster — lenses exist, profile fields empty. A prose
    # markdown roster is bigram_fallback evidence, so the teacher confirms
    # the two names (STEP 3) before anything is created.
    roster = "Class list\n\nMarco Bianchi\n\nNora Rossi\n"
    job = _run_to_done(
        _upload_bytes("roster.md", roster.encode("utf-8"), "text/markdown")["job_id"]
    )
    assert job["status"] == "done"
    job = _confirm_pending(job)
    ids_before = {s["display_name"]: s["student_id"] for s in job["students_created"]}
    assert sorted(ids_before) == ["Marco Bianchi", "Nora Rossi"]

    # Day 30: support document arrives, mentioning both by name.
    support_doc = (
        "Learning support notes\n\n"
        "Marco Bianchi benefits from a checklist before writing.\n\n"
        "Nora Rossi has strong inference skills and can extend ideas with a relevant quotation.\n"
    )
    job2 = _run_to_done(
        _upload_bytes("support-notes.md", support_doc.encode("utf-8"), "text/markdown")["job_id"]
    )
    assert job2["status"] == "done"
    job2 = _confirm_pending(job2)
    updated = {s["display_name"]: s for s in job2["students_created"]}
    # Same student_ids: the import merged, it did not mint new students.
    for name, sid in ids_before.items():
        assert updated[name]["student_id"] == sid

    # The roster still holds exactly two students.
    names = web._with_student_store(
        lambda store: sorted(lens["display_name"] for lens in store.list_lenses())
    )
    assert names == ["Marco Bianchi", "Nora Rossi"]

    # And the existing lenses gained evidenced fields from the document.
    assert "executive_functioning" in updated["Marco Bianchi"]["fields_populated"]
    assert "academic_strengths" in updated["Nora Rossi"]["fields_populated"]
    root = vault.vault_root()
    lens = json.loads(
        (root / "lenses" / ids_before["Marco Bianchi"] / "lens.json").read_text(encoding="utf-8")
    )
    field = lens["profile"]["executive_functioning"]
    assert field["value"], field
    assert field["evidence"], field


def test_drive_folder_link_routes_to_class_folder_ingest(isolated_state, monkeypatch):
    """Operator request 2026-08-18: pasting a Drive FOLDER link into the
    Students import box runs the class-folder ingest (match every document
    inside by student name) instead of failing as a bad file link."""
    from src.lingua_viva import class_folder_ingest

    seen = {}

    def fake_ingest(folder_id, teacher_id, *, store, **kwargs):
        seen["folder_id"] = folder_id
        return {
            "folder_id": folder_id,
            "teacher_id": teacher_id,
            "files_seen": 3,
            "files_processed": 3,
            "students_created_or_updated": [
                {"student_id": "student-marco-bianchi", "display_name": "Marco Bianchi",
                 "fields_populated": ["executive_functioning"]},
            ],
            "unattributed": [],
            "failed": [],
            "truncated": False,
        }

    monkeypatch.setattr(class_folder_ingest, "ingest_class_folder", fake_ingest)
    response = client.post(
        "/api/students/ingest",
        json={"drive_ref": "https://drive.google.com/drive/folders/AbCdEf12345?usp=sharing"},
    )
    assert response.status_code == 200, response.text
    data = response.json()
    assert data["kind"] == "class_folder"
    assert seen["folder_id"] == "AbCdEf12345"
    assert data["students_created_or_updated"][0]["display_name"] == "Marco Bianchi"


def test_drive_folder_link_without_connection_says_connect_first(isolated_state):
    response = client.post(
        "/api/students/ingest",
        json={"drive_ref": "https://drive.google.com/drive/folders/AbCdEf12345"},
    )
    assert response.status_code == 409
    assert "Connect Google Drive first" in response.json()["error"]


def test_drive_folder_link_without_readable_id_is_400(isolated_state):
    response = client.post(
        "/api/students/ingest",
        json={"drive_ref": "https://drive.google.com/drive/folders/"},
    )
    assert response.status_code == 400
    assert "folder" in response.json()["error"].lower()


def test_real_unsupported_upload_fails_honestly(isolated_state, monkeypatch):
    import src.lingua_viva.docpipe.model as docpipe_model

    monkeypatch.setattr(
        docpipe_model, "LocalModelClient",
        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("no model in test")),
    )
    started = _upload_bytes(
        "slides.pptx",
        b"not really a presentation, but the extension is unsupported",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    )
    job = _run_to_done(started["job_id"])

    assert job["status"] == "failed"
    assert "unsupported format" in job["error"]
    assert "xlsx" in job["error"] and "docx" in job["error"]


def test_no_students_found_is_honest(isolated_state, monkeypatch):
    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        data["structure"]["students_detected"] = []
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _run_to_done(_upload()["job_id"])
    assert job["status"] == "done"
    assert job["students_found"] == 0
    assert job["students_created"] == []


def test_oversized_upload_is_rejected(isolated_state, monkeypatch):
    monkeypatch.setattr(web, "INGEST_MAX_BYTES", 10)
    response = client.post(
        "/api/students/ingest",
        files={"file": ("big.md", b"x" * 11, "text/markdown")},
    )
    assert response.status_code == 413


def test_unknown_job_says_reimport_is_safe(isolated_state):
    response = client.get("/api/students/ingest/JOB-nope")
    assert response.status_code == 404
    assert "safe" in response.json()["error"]


def test_drive_ref_without_drive_connection_says_connect_first(isolated_state):
    """G2: fetch_file is real now. With no Drive connection the endpoint asks
    the teacher to connect first (409) instead of pretending to work."""
    response = client.post("/api/students/ingest", json={"drive_ref": "folder-123"})
    assert response.status_code == 409
    assert "Connect Google Drive first" in response.json()["error"]


def test_drive_ref_unparseable_link_is_a_clear_400(isolated_state):
    response = client.post("/api/students/ingest", json={"drive_ref": "???"})
    assert response.status_code == 400
    assert "Drive file link" in response.json()["error"]


def test_no_demo_roster_seeding_remains():
    source = (REPO / "src" / "web.py").read_text(encoding="utf-8")
    assert "_seed_demo_roster" not in source
    assert 'student_id="student-luca"' not in source


HTML = (REPO / "static" / "index.html").read_text(encoding="utf-8")


def test_import_affordance_present_no_manual_entry():
    assert "Import your roster" in HTML
    assert 'id="ingest-file"' in HTML
    assert '"/api/students/ingest"' in HTML
    assert "No student names found" in HTML
    assert ".xlsx,.docx,.csv,.txt,.md,.markdown,.pdf" in HTML


def test_drive_roster_import_affordance_wired():
    """G2 UI: the Students view offers Drive import — connect when possible,
    paste-link import when connected, honest message otherwise."""
    assert 'id="drive-roster-import"' in HTML
    assert 'id="drive-roster-link"' in HTML
    assert "Import from Drive" in HTML
    assert "Connect Google Drive" in HTML
    assert "drive_ref: link" in HTML
    assert '"/api/google-drive/auth/start"' in HTML


def test_drive_folder_import_affordance_wired():
    """Operator request 2026-08-18: the same paste-link box accepts a folder
    of student documents, and the UI renders the folder summary."""
    assert "or a folder of student documents" in HTML
    assert 'kind === "class_folder"' in HTML
    assert "students_created_or_updated" in HTML


def test_per_student_remove_control_wired():
    """Operator request 2026-08-18: removing one student (left the school,
    teacher imported as a student) must be possible from the roster itself,
    not only from the bottom of an opened lens."""
    assert "data-student-remove" in HTML
    assert "data-student-remove-name" in HTML
    assert "Remove ${name} from your roster?" in HTML


def test_low_confidence_confirm_button_wired():
    assert "data-ingest-confirm" in HTML
    assert '"/api/students/ingest/confirm"' in HTML


def test_bulk_roster_confirm_and_undo_controls_wired():
    assert "data-ingest-confirm-check" in HTML
    assert "data-ingest-confirm-selected" in HTML
    assert "display_names: names" in HTML
    assert "Undo this import" in HTML
    assert "/api/students/ingest/${state.ingestJob.job_id}" in HTML
    assert '{method: "DELETE"}' in HTML


def test_ingest_teacher_id_passthrough_registers_roster(isolated_state, fixture_extractor):
    """Prepare-fix P3b: the teacher who imports a document owns the created
    students on their roster (teacher_roster, source='ingest')."""
    content = (FIXTURES / "lesson_plan_marco_nora.md").read_bytes()
    response = client.post(
        "/api/students/ingest",
        files={"file": ("lesson_plan_marco_nora.md", content, "text/markdown")},
        data={"teacher_id": "teacher-claudia"},
    )
    assert response.status_code == 200, response.text
    job = _run_to_done(response.json()["job_id"])
    assert job["status"] == "done", job

    def roster_rows(store):
        rows = store._conn.execute(
            "SELECT student_id, source FROM teacher_roster WHERE teacher_id = ?",
            ("teacher-claudia",),
        ).fetchall()
        return sorted((r["student_id"], r["source"]) for r in rows)

    rows = web._with_student_store(roster_rows)
    assert rows, "ingest did not register roster membership"
    assert all(source == "ingest" for _sid, source in rows)

    def scoped(store):
        return sorted(str(l.get("display_name")) for l in store.list_lenses_for_teacher("teacher-claudia"))

    assert web._with_student_store(scoped) == ["Marco Bianchi", "Nora Rossi"]


def test_confirm_accepts_optional_cefr_level(isolated_state, monkeypatch):
    """Prepare-fix P3c: the confirmation step can carry a starting CEFR
    level per student; it lands in the snapshot via audited observations."""
    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        for student in data["structure"]["students_detected"]:
            student["evidence"] = "bigram_fallback"
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _run_to_done(_upload()["job_id"])
    names = [item["display_name"] for item in job["needs_confirmation"]]
    assert "Nora Rossi" in names

    response = client.post("/api/students/ingest/confirm", json={
        "job_id": job["job_id"],
        "display_names": names,
        "cefr_levels": {"Nora Rossi": "A2"},
    })
    assert response.status_code == 200, response.text
    by_name = {s["display_name"]: s for s in response.json()["students"]}
    assert by_name["Nora Rossi"]["cefr_level"] == "A2"
    assert "cefr_level" not in by_name["Marco Bianchi"]

    def snapshot(store):
        return (
            store.get_lens(by_name["Nora Rossi"]["student_id"])["cefr_snapshot"],
            store.get_lens(by_name["Marco Bianchi"]["student_id"])["cefr_snapshot"],
        )

    nora, marco = web._with_student_store(snapshot)
    assert nora == {"reading": "A2", "writing": "A2", "speaking": "A2", "listening": "A2"}
    assert set(marco.values()) == {None}


def test_confirm_rejects_invalid_cefr_level(isolated_state, monkeypatch):
    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        for student in data["structure"]["students_detected"]:
            student["evidence"] = "bigram_fallback"
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _run_to_done(_upload()["job_id"])

    response = client.post("/api/students/ingest/confirm", json={
        "job_id": job["job_id"],
        "display_name": "Nora Rossi",
        "cefr_level": "Z9",
    })
    assert response.status_code == 400
    assert "invalid_cefr_level" in response.json()["error"]
    # Nothing was created by the rejected request.
    refreshed = client.get(f"/api/students/ingest/{job['job_id']}").json()
    assert len(refreshed["needs_confirmation"]) == 2

def test_confirm_sets_levels_for_auto_created_students(isolated_state, fixture_extractor):
    """Claudia-lens audit P0-1 (2026-08-18): roster imports auto-create every
    student, so the confirm route must also accept starting levels for
    already-created students — with no display_names at all."""
    job = _run_to_done(_upload()["job_id"])
    assert job["status"] == "done", job
    created = {s["display_name"]: s for s in job["students_created"]}
    assert "Nora Rossi" in created and "Marco Bianchi" in created

    response = client.post("/api/students/ingest/confirm", json={
        "job_id": job["job_id"],
        "cefr_levels": {"Nora Rossi": "B1"},
    })
    assert response.status_code == 200, response.text
    body = response.json()
    assert body["status"] == "updated"
    assert body["students"] == []
    updated = {u["display_name"]: u for u in body["cefr_updated"]}
    assert updated["Nora Rossi"]["cefr_level"] == "B1"

    def snapshots(store):
        return (
            store.get_lens(created["Nora Rossi"]["student_id"])["cefr_snapshot"],
            store.get_lens(created["Marco Bianchi"]["student_id"])["cefr_snapshot"],
        )

    nora, marco = web._with_student_store(snapshots)
    assert nora == {"reading": "B1", "writing": "B1", "speaking": "B1", "listening": "B1"}
    assert set(marco.values()) == {None}

    # The level persists on the job record so re-renders show it as saved.
    refreshed = client.get(f"/api/students/ingest/{job['job_id']}").json()
    by_name = {s["display_name"]: s for s in refreshed["students_created"]}
    assert by_name["Nora Rossi"]["cefr_level"] == "B1"


def test_created_row_cefr_controls_wired():
    """Claudia-lens audit P0-1: the starting-level dropdown renders on
    auto-created rows too, with its own save control."""
    assert "data-created-cefr" in HTML
    assert "data-ingest-save-levels" in HTML
    assert "Save starting levels" in HTML


# ---------------------------------------------------------------------------
# Phase 0A — always-preview ingest (SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19
# §2A). The class lock: preview writes NOTHING — no student store rows, no
# vault lenses, no Drive sync enqueue. Only the explicit approve creates.
# ---------------------------------------------------------------------------

_ROSTER_NAMES = ["Marco Bianchi", "Nora Rossi", "Luca Verdi", "Sara Conti", "Maya Singh"]


@pytest.fixture()
def roster_extractor(monkeypatch):
    async def extract(source, content, *, model_client=None):
        data = _roster_extraction(source.source_id, _ROSTER_NAMES)
        data["source_sha256"] = source.data["sha256"]
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)


@pytest.fixture()
def sync_spies(monkeypatch):
    """Spy on every surface that could push a student toward Drive."""
    from src.lingua_viva import drive_sync
    from src.lingua_viva.docpipe import sync as docpipe_sync

    calls = {"enqueued": [], "folder": 0, "triggered": []}
    monkeypatch.setattr(docpipe_sync, "enqueue_lens", lambda sid: calls["enqueued"].append(sid))
    monkeypatch.setattr(drive_sync, "ensure_lens_sync_folder",
                        lambda: calls.__setitem__("folder", calls["folder"] + 1))
    monkeypatch.setattr(drive_sync, "trigger_sync", lambda sid: calls["triggered"].append(sid))
    return calls


def _store_count() -> int:
    return web._with_student_store(lambda store: len(store.list_lenses()))


def test_preview_writes_nothing(isolated_state, roster_extractor, sync_spies):
    """The Phase 0A class lock. Store count identical before and after the
    preview; nothing enqueued toward Drive; the full would-be result is
    reported per student."""
    assert _store_count() == 0
    job = _wait_for_job(_upload()["job_id"])

    assert job["status"] == "preview"
    assert [s["display_name"] for s in job["preview_students"]] == _ROSTER_NAMES
    for student in job["preview_students"]:
        assert student["span_ids"], student
        assert student["low_confidence"] is False
    assert job["students_created"] == []
    assert job["needs_confirmation"] == []

    # Nothing was written anywhere.
    assert _store_count() == 0
    assert sync_spies == {"enqueued": [], "folder": 0, "triggered": []}
    lenses_dir = vault.vault_root() / "lenses"
    assert not lenses_dir.exists() or not any(lenses_dir.iterdir())


def test_preview_flags_low_confidence_names_without_creating(isolated_state, monkeypatch, sync_spies):
    async def extract(source, content, *, model_client=None):
        data = _roster_extraction(source.source_id, _ROSTER_NAMES)
        data["source_sha256"] = source.data["sha256"]
        for student in data["structure"]["students_detected"]:
            if student["display_name"] == "Maya Singh":
                student["evidence"] = "bigram_fallback"
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _wait_for_job(_upload()["job_id"])

    assert job["status"] == "preview"
    flags = {s["display_name"]: s["low_confidence"] for s in job["preview_students"]}
    assert flags["Maya Singh"] is True
    assert flags["Marco Bianchi"] is False
    assert _store_count() == 0


def test_approve_creates_and_syncs(isolated_state, roster_extractor, sync_spies):
    job = _wait_for_job(_upload()["job_id"])
    assert job["status"] == "preview"

    response = client.post("/api/students/ingest/approve", json={"job_id": job["job_id"]})
    assert response.status_code == 200, response.text
    assert response.json()["status"] == "creating"

    done = _wait_for_job(job["job_id"])
    assert done["status"] == "done"
    assert [s["display_name"] for s in done["students_created"]] == _ROSTER_NAMES
    assert _store_count() == len(_ROSTER_NAMES)
    created_ids = sorted(s["student_id"] for s in done["students_created"])
    assert sorted(sync_spies["enqueued"]) == created_ids
    assert sync_spies["folder"] == 1
    assert sorted(sync_spies["triggered"]) == created_ids


def test_cancel_leaves_zero_trace(isolated_state, roster_extractor, sync_spies):
    job = _wait_for_job(_upload()["job_id"])
    assert job["status"] == "preview"

    response = client.post("/api/students/ingest/cancel", json={"job_id": job["job_id"]})
    assert response.status_code == 200, response.text
    assert response.json()["status"] == "cancelled"

    assert _store_count() == 0
    assert sync_spies == {"enqueued": [], "folder": 0, "triggered": []}
    refreshed = client.get(f"/api/students/ingest/{job['job_id']}").json()
    assert refreshed["status"] == "cancelled"
    assert refreshed["students_created"] == []

    # A cancelled preview can never be approved afterwards.
    response = client.post("/api/students/ingest/approve", json={"job_id": job["job_id"]})
    assert response.status_code == 409
    assert _store_count() == 0


def test_double_approve_is_refused(isolated_state, roster_extractor, sync_spies):
    job = _wait_for_job(_upload()["job_id"])
    first = client.post("/api/students/ingest/approve", json={"job_id": job["job_id"]})
    assert first.status_code == 200, first.text
    done = _wait_for_job(job["job_id"])
    assert done["status"] == "done"

    second = client.post("/api/students/ingest/approve", json={"job_id": job["job_id"]})
    assert second.status_code == 409
    # No duplicates from the refused second approve.
    assert _store_count() == len(_ROSTER_NAMES)


def test_cancel_requires_a_preview(isolated_state, roster_extractor):
    job = _run_to_done(_upload()["job_id"])
    assert job["status"] == "done"
    response = client.post("/api/students/ingest/cancel", json={"job_id": job["job_id"]})
    assert response.status_code == 409


def test_approve_and_cancel_unknown_job_is_404(isolated_state):
    assert client.post("/api/students/ingest/approve", json={"job_id": "JOB-nope"}).status_code == 404
    assert client.post("/api/students/ingest/cancel", json={"job_id": "JOB-nope"}).status_code == 404


def test_zero_detections_skip_the_preview(isolated_state, monkeypatch):
    """Nothing to approve → the job finishes honestly as done with zero
    creations instead of asking the teacher to confirm an empty list."""
    async def extract(source, content, *, model_client=None):
        data = _fixture_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        data["structure"]["students_detected"] = []
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)
    job = _wait_for_job(_upload()["job_id"])
    assert job["status"] == "done"
    assert job["students_found"] == 0
    assert job["preview_students"] == []
    assert job["students_created"] == []


def test_preview_survives_memory_reset_and_still_approves(isolated_state, roster_extractor):
    """The preview is durable: an app restart between preview and approve
    must not lose the teacher's pending import."""
    job = _wait_for_job(_upload()["job_id"])
    assert job["status"] == "preview"
    web._INGEST_JOBS = {}

    response = client.post("/api/students/ingest/approve", json={"job_id": job["job_id"]})
    assert response.status_code == 200, response.text
    done = _wait_for_job(job["job_id"])
    assert done["status"] == "done"
    assert [s["display_name"] for s in done["students_created"]] == _ROSTER_NAMES


def test_preview_controls_wired():
    """Phase 0A UI: the preview panel offers explicit create/cancel and says
    honestly that nothing is saved yet."""
    assert "data-ingest-approve" in HTML
    assert "data-ingest-cancel" in HTML
    assert '"/api/students/ingest/approve"' in HTML
    assert '"/api/students/ingest/cancel"' in HTML
    assert "Create these ${scoped.length} student" in HTML
    assert "Nothing has been saved yet" in HTML
    # Claudia-lens rule (v157): raw confidence numbers never render in the UI.
    assert "student.confidence" not in HTML


# --- STEP 4 (SPEC_LV_UNIFIED_REAL_DATA_FIX §STEP 4): "only my class" -----------


_CLASS_LIST = {
    "Verdi": {
        "teacher": "Ilaria Moretti",
        "students": ["Marco Bianchi", "Nora Rossi", "Sara Conti"],
    },
    "Blu": {
        "teacher": "Chiara De Luca",
        "students": ["Matteo Barbieri", "Sofia Caruso"],
    },
}


def _class_list_extraction(source_id: str) -> dict:
    """Extraction shaped like the class-pair detector's output: every
    detection carries class / grade / teacher_attribution (STEP 4)."""
    data = _roster_extraction(
        source_id,
        [name for info in _CLASS_LIST.values() for name in info["students"]],
    )
    by_name = {
        name: (class_name, info["teacher"])
        for class_name, info in _CLASS_LIST.items()
        for name in info["students"]
    }
    for student in data["structure"]["students_detected"]:
        class_name, teacher = by_name[student["display_name"]]
        student["class"] = class_name
        student["grade"] = "Grade 3"
        student["teacher_attribution"] = teacher
    return data


@pytest.fixture()
def class_list_extractor(monkeypatch):
    async def extract(source, content, *, model_client=None):
        data = _class_list_extraction(source.source_id)
        data["source_sha256"] = source.data["sha256"]
        return ExtractionRecord(data)

    from src.lingua_viva.docpipe import extract as docpipe_extract

    monkeypatch.setattr(docpipe_extract, "extract_document", extract)


def test_preview_carries_class_membership(isolated_state, class_list_extractor, sync_spies):
    """Preview rows surface class / grade / teacher_attribution and the job
    lists the distinct classes so the UI can offer the scope picker."""
    job = _wait_for_job(_upload()["job_id"])
    assert job["status"] == "preview"
    assert job["preview_classes"] == ["Blu", "Verdi"]
    rows = {row["display_name"]: row for row in job["preview_students"]}
    assert rows["Marco Bianchi"]["class"] == "Verdi"
    assert rows["Marco Bianchi"]["grade"] == "Grade 3"
    assert rows["Marco Bianchi"]["teacher_attribution"] == "Ilaria Moretti"
    assert rows["Sofia Caruso"]["class"] == "Blu"
    assert rows["Sofia Caruso"]["teacher_attribution"] == "Chiara De Luca"
    assert _store_count() == 0  # still a preview — nothing written


def test_approve_scoped_to_one_class_creates_only_that_class(isolated_state, class_list_extractor, sync_spies):
    """The product requirement: the teacher wants her ~39, not 400. A scoped
    approve creates only her class and says honestly how many it skipped."""
    job = _wait_for_job(_upload()["job_id"])
    assert job["status"] == "preview"

    response = client.post("/api/students/ingest/approve", json={
        "job_id": job["job_id"], "classes": ["Verdi"],
    })
    assert response.status_code == 200, response.text
    done = _wait_for_job(job["job_id"])
    assert done["status"] == "done"
    created = sorted(s["display_name"] for s in done["students_created"])
    assert created == sorted(_CLASS_LIST["Verdi"]["students"])
    assert _store_count() == 3
    assert any("Skipped 2" in w and "Verdi" in w for w in done["warnings"]), done["warnings"]
    # only the created students head toward Drive
    assert len(sync_spies["enqueued"]) == 3


def test_unscoped_approve_creates_every_class(isolated_state, class_list_extractor, sync_spies):
    job = _wait_for_job(_upload()["job_id"])
    response = client.post("/api/students/ingest/approve", json={"job_id": job["job_id"]})
    assert response.status_code == 200, response.text
    done = _wait_for_job(job["job_id"])
    assert done["status"] == "done"
    assert _store_count() == 5
    assert not any("Skipped" in w for w in done["warnings"])


def test_approve_with_unknown_class_is_refused(isolated_state, class_list_extractor, sync_spies):
    """An unknown class selection is a 422, never a silent zero-creation
    approve — the job stays at preview and can still be approved correctly."""
    job = _wait_for_job(_upload()["job_id"])
    response = client.post("/api/students/ingest/approve", json={
        "job_id": job["job_id"], "classes": ["Rossi"],
    })
    assert response.status_code == 422
    assert _store_count() == 0
    refreshed = client.get(f"/api/students/ingest/{job['job_id']}").json()
    assert refreshed["status"] == "preview"

    # a correct scope still works afterwards
    response = client.post("/api/students/ingest/approve", json={
        "job_id": job["job_id"], "classes": ["Blu"],
    })
    assert response.status_code == 200, response.text
    done = _wait_for_job(job["job_id"])
    assert sorted(s["display_name"] for s in done["students_created"]) == sorted(
        _CLASS_LIST["Blu"]["students"]
    )


def test_class_scope_ui_wired():
    """STEP 4 UI: the preview offers the class scope picker and per-row class
    badges; the approve sends the selected class."""
    assert "data-ingest-class-scope" in HTML
    assert "preview_classes" in HTML
    assert "you can import only yours" in HTML
    assert "classes: [scope]" in HTML
