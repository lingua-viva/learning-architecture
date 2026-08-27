"""T3 — grounded extraction (SPEC_T3_EXTRACTION_2026-08-04).

The forced-hallucination test is the point of the workstream: a model claim
whose cited span does not support it must be DROPPED, never demoted.
"""
from __future__ import annotations

import asyncio
import json
from pathlib import Path

import pytest

from src.lingua_viva.docpipe import jobs, vault
from src.lingua_viva.docpipe.contracts import SourceRecord
from src.lingua_viva.docpipe.extract import extract_document
from src.lingua_viva.docpipe.grounding_docs import verify_extraction
from src.lingua_viva.docpipe.model import ModelResult

REPO = Path(__file__).resolve().parents[1]
FIXTURES = REPO / "tests" / "fixtures" / "docpipe"


def _source(fixture_source: str, fixture_content: str) -> tuple[SourceRecord, bytes]:
    data = json.loads((FIXTURES / fixture_source).read_text(encoding="utf-8"))
    content = (FIXTURES / fixture_content).read_bytes()
    return SourceRecord(data), content


def _run(coro):
    return asyncio.get_event_loop_policy().new_event_loop().run_until_complete(coro)


# --- Fixture parity (grounding pass rate must be 100%) -------------------------


def test_lesson_plan_fixture_parity():
    expected = json.loads(
        (FIXTURES / "expected_extraction_lesson_plan_marco_nora.json").read_text(encoding="utf-8")
    )
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    record = _run(extract_document(source, content))
    data = record.data

    assert data["normalized_text"] == expected["normalized_text"]
    assert data["spans"] == expected["spans"]
    assert data["structure"]["title"] == expected["structure"]["title"]
    assert data["structure"]["document_type"] == expected["structure"]["document_type"]
    assert data["structure"]["sections"] == expected["structure"]["sections"]
    assert data["structure"]["students_detected"] == expected["structure"]["students_detected"]
    assert data["structure"]["curriculum"] == expected["structure"]["curriculum"]
    assert data["language"] == expected["language"]
    assert data["source_sha256"] == source.data["sha256"]

    report = verify_extraction(data)
    assert report.ok and not report.dropped, f"grounding pass rate below 100%: {report}"


def test_student_work_fixture_parity():
    expected = json.loads(
        (FIXTURES / "expected_extraction_student_work_nora_rossi.json").read_text(encoding="utf-8")
    )
    source, content = _source("source_student_work_nora_rossi.json", "student_work_nora_rossi.md")
    record = _run(extract_document(source, content))
    data = record.data

    assert data["normalized_text"] == expected["normalized_text"]
    assert data["spans"] == expected["spans"]
    assert data["structure"]["title"] == expected["structure"]["title"]
    assert data["structure"]["document_type"] == expected["structure"]["document_type"]
    assert data["structure"]["sections"] == expected["structure"]["sections"]
    assert data["structure"]["students_detected"] == expected["structure"]["students_detected"]
    assert data["structure"]["curriculum"] == expected["structure"]["curriculum"]

    report = verify_extraction(data)
    assert report.ok and not report.dropped


def test_extraction_is_schema_valid():
    from src.lingua_viva.docpipe.validate import validate_file

    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    record = _run(extract_document(source, content))
    out = FIXTURES.parent / ".." / ".tmp-extract-check.json"
    out = out.resolve()
    out.write_text(json.dumps(record.data), encoding="utf-8")
    try:
        # validator infers schema by filename shape; write into an extracted/ shape
        pass
    finally:
        out.unlink(missing_ok=True)
    # direct schema check through the vault write path instead:
    import tempfile

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        vault.put_source(source, content, root=root)
        vault.put_extraction(record, root=root)  # raises if schema-invalid
        assert (root / "extracted" / f"{source.source_id}.json").exists()


# --- THE hallucination test ----------------------------------------------------


class ScriptedModel:
    def __init__(self, replies):
        self.replies = list(replies)
        self.calls = 0

    async def complete(self, prompt, *, system_prompt=None, context=None, max_tokens=2000):
        self.calls += 1
        reply = self.replies.pop(0) if self.replies else self.replies_default
        return ModelResult(content=reply, confidence=0.9, model_used="scripted")


def test_hallucinated_student_is_dropped_not_demoted():
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    model = ScriptedModel([json.dumps({
        "students": [
            {"display_name": "Giulia Ferrari", "span_id": "SPN-0003"},   # span has no such name
            {"display_name": "Luca Verdi", "span_id": "SPN-9999"},       # span does not exist
        ]
    })])
    record = _run(extract_document(source, content, model_client=model))
    names = [s["display_name"] for s in record.data["structure"]["students_detected"]]
    assert "Giulia Ferrari" not in names
    assert "Luca Verdi" not in names
    assert names == ["Marco Bianchi", "Nora Rossi"]  # deterministic finds intact
    dropped = [w for w in record.data["warnings"] if w.startswith("grounding_dropped:model_student:")]
    assert len(dropped) == 2
    # No lowered-confidence survivors: every remaining entry is fully supported.
    assert verify_extraction(record.data).ok


def test_supported_model_student_is_added_with_lower_confidence():
    source, content = _source("source_student_work_nora_rossi.json", "student_work_nora_rossi.md")
    # SPN-0004 really contains no student name; SPN-0003 contains "Nora" (already found).
    # Give the model a claim that IS supported: cite SPN-0001 for Nora Rossi (dupe → ignored),
    # then verify a genuinely-new supported name in a synthetic doc instead.
    # Lowercase in the document → invisible to the capitalized-bigram
    # detector; only the model can claim it, and the claim IS span-supported
    # (verification is case/accent-folded).
    synthetic = (
        "# Class list\n\nRoster: G3 Italian\n\nnew arrival: ada colombo\n\nMarco Bianchi\n\n"
    ).encode("utf-8")
    src = SourceRecord({**json.loads((FIXTURES / "source_student_work_nora_rossi.json").read_text()),
                        "sha256": __import__("hashlib").sha256(synthetic).hexdigest()})
    model = ScriptedModel([json.dumps({
        "students": [{"display_name": "Ada Colombo", "span_id": "SPN-0003"}]
    })])
    record = _run(extract_document(src, synthetic, model_client=model))
    by_name = {s["display_name"]: s for s in record.data["structure"]["students_detected"]}
    assert "Ada Colombo" in by_name
    assert by_name["Ada Colombo"]["confidence"] == 0.7
    assert verify_extraction(record.data).ok


def test_malformed_model_json_retries_then_degrades_honestly():
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    model = ScriptedModel(["I think the students are Marco and Nora!", "still not json"])
    record = _run(extract_document(source, content, model_client=model))
    assert model.calls == 2
    assert any(w.startswith("model_enrichment_discarded") for w in record.data["warnings"])
    assert [s["display_name"] for s in record.data["structure"]["students_detected"]] == [
        "Marco Bianchi", "Nora Rossi",
    ]


# --- STEP 6 (SPEC_LV_UNIFIED_REAL_DATA_FIX §STEP 6): the enrichment veto -------
#
# Additive-only enrichment cannot converge on truth. The model may now
# DISPUTE a detection — but a veto never deletes: grounded claims set
# removal_proposed for downstream review; ungrounded claims are DROPPED
# with the same mechanical severity as ungrounded additions.


def _span_containing(record, text: str) -> str:
    return next(s["span_id"] for s in record.data["spans"] if text in str(s["text"]))


def test_model_veto_flags_never_deletes():
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    plain = _run(extract_document(source, content))
    span_id = _span_containing(plain, "Nora Rossi")
    model = ScriptedModel([json.dumps({
        "students": [],
        "not_students": [{
            "display_name": "Nora Rossi", "span_id": span_id,
            "reason": "listed as the classroom tutor",
        }],
    })])
    record = _run(extract_document(source, content, model_client=model))
    by_name = {s["display_name"]: s for s in record.data["structure"]["students_detected"]}
    # the veto NEVER deletes — the detection stays, flagged for review
    assert "Nora Rossi" in by_name
    assert by_name["Nora Rossi"]["removal_proposed"] == "listed as the classroom tutor"
    assert "removal_proposed" not in by_name["Marco Bianchi"]
    assert any(w.startswith("model_enrichment_veto:Nora Rossi") for w in record.data["warnings"])
    assert verify_extraction(record.data).ok

    # the flagged extraction survives the frozen vault schema gate
    import tempfile

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        vault.put_source(source, content, root=root)
        vault.put_extraction(record, root=root)


def test_ungrounded_vetoes_are_dropped_like_ungrounded_additions():
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    plain = _run(extract_document(source, content))
    nora_span = _span_containing(plain, "Nora Rossi")
    no_name_span = next(
        s["span_id"] for s in plain.data["spans"]
        if "Nora" not in str(s["text"]) and "Marco" not in str(s["text"])
    )
    model = ScriptedModel([json.dumps({
        "students": [],
        "not_students": [
            # not a current detection
            {"display_name": "Giulia Ferrari", "span_id": nora_span, "reason": "a teacher"},
            # span does not exist
            {"display_name": "Nora Rossi", "span_id": "SPN-9999", "reason": "a teacher"},
            # no reason given
            {"display_name": "Nora Rossi", "span_id": nora_span, "reason": ""},
            # cited span does not contain the name
            {"display_name": "Marco Bianchi", "span_id": no_name_span, "reason": "a teacher"},
        ],
    })])
    record = _run(extract_document(source, content, model_client=model))
    students = record.data["structure"]["students_detected"]
    assert all("removal_proposed" not in s for s in students)
    assert [s["display_name"] for s in students] == ["Marco Bianchi", "Nora Rossi"]
    dropped = [w for w in record.data["warnings"] if w.startswith("grounding_dropped:model_veto:")]
    assert len(dropped) == 4


def test_veto_matching_uses_the_one_normalizer():
    """§5 prohibition: ONE normalizer. A case-variant veto still matches the
    detection through identity.normalize_name — no second comparison path."""
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    plain = _run(extract_document(source, content))
    span_id = _span_containing(plain, "Nora Rossi")
    model = ScriptedModel([json.dumps({
        "students": [],
        "not_students": [{
            "display_name": "NORA  ROSSI", "span_id": span_id, "reason": "a tutor",
        }],
    })])
    record = _run(extract_document(source, content, model_client=model))
    by_name = {s["display_name"]: s for s in record.data["structure"]["students_detected"]}
    assert by_name["Nora Rossi"]["removal_proposed"] == "a tutor"


# --- Span integrity + offline + formats ---------------------------------------


def test_spans_slice_back_exactly():
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    record = _run(extract_document(source, content))
    text = record.data["normalized_text"]
    for span in record.data["spans"]:
        assert span["text"] == text[span["char_start"]:span["char_end"]]


def test_tampered_span_fails_verification():
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    record = _run(extract_document(source, content))
    record.data["spans"][0]["char_end"] += 3
    report = verify_extraction(record.data)
    assert not report.ok
    assert any("span_integrity" in e for e in report.errors)


def test_offline_no_model_is_fully_deterministic_with_warning():
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    record = _run(extract_document(source, content, model_client=None))
    assert record.data["extractor"]["model"] is None
    assert any(w.startswith("model_enrichment_unavailable") for w in record.data["warnings"])
    assert len(record.data["structure"]["students_detected"]) == 2


# --- STEP 7 (SPEC_LV_UNIFIED_REAL_DATA_FIX §STEP 7, L6): honest failure reporting


def test_privacy_refusal_reports_true_reason_not_invalid_json():
    """A none:local_only refusal must surface as model_enrichment_unavailable
    with the privacy reason. The 08-19 audit found this exact case
    misreported as 'invalid JSON after retry' — the model never ran at all."""

    class RefusingModel:
        async def complete(self, prompt, *, system_prompt=None, context=None, max_tokens=2000):
            return ModelResult(
                content="Student data must stay on this computer and no local model is installed.",
                confidence=0.0,
                model_used="none:local_only",
                error="local_only_no_model",
            )

    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    record = _run(extract_document(source, content, model_client=RefusingModel()))
    warnings = record.data["warnings"]
    assert any(w == "model_enrichment_unavailable:local_only_no_model" for w in warnings)
    assert not any("invalid JSON" in w for w in warnings)
    # deterministic detection is untouched by the refusal
    assert [s["display_name"] for s in record.data["structure"]["students_detected"]] == [
        "Marco Bianchi", "Nora Rossi",
    ]


def test_none_model_without_error_still_reports_unavailable():
    """Belt for the same class: even when a refusal arrives with error=''
    (the pre-fix engine shape), a 'none*' model_used means the model never
    ran — its prose must never reach the JSON parser and be blamed as
    invalid JSON."""

    class ProseRefusal:
        async def complete(self, prompt, *, system_prompt=None, context=None, max_tokens=2000):
            return ModelResult(
                content="Install a local model first.",
                confidence=0.0,
                model_used="none:local_only",
            )

    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    record = _run(extract_document(source, content, model_client=ProseRefusal()))
    warnings = record.data["warnings"]
    assert any(w.startswith("model_enrichment_unavailable:no usable local model") for w in warnings)
    assert not any("invalid JSON" in w for w in warnings)


def test_local_model_client_respects_engine_resolution_order():
    """STEP 8: LocalModelClient must not hard-code detect_model() — passing
    an explicit model jumped the engine's resolution order, so the
    LV_REASON_MODEL override was silently ignored on every docpipe call."""
    from src.lingua_viva.docpipe.model import LocalModelClient
    from src.lingua_viva.reasoning import ReasonResult

    captured: dict = {}

    class Engine:
        async def reason(self, prompt, **kwargs):
            captured.update(kwargs)
            return ReasonResult(content="{}", confidence=0.75, model_used="ollama/qwen2.5:7b")

    result = _run(LocalModelClient(engine=Engine()).complete("prompt", system_prompt="sys"))

    assert "model" not in captured  # the engine resolves; overrides work
    assert captured["local_only"] is True
    assert result.model_used == "ollama/qwen2.5:7b"


def test_unsupported_format_fails_honestly():
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    bad = SourceRecord({**source.data, "mime": "image/png", "original_ext": ".png"})
    with pytest.raises(ValueError, match="unsupported format"):
        _run(extract_document(bad, b"\x89PNG..."))


def test_xlsx_roster_extracts_cell_text():
    from io import BytesIO

    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["First", "Last"])
    sheet.append(["Marco", "Bianchi"])
    sheet.append(["Nora", "Rossi"])
    buffer = BytesIO()
    workbook.save(buffer)
    source, _ = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    xlsx = SourceRecord({
        **source.data,
        "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "original_ext": ".xlsx",
    })

    record = _run(extract_document(xlsx, buffer.getvalue()))

    assert "Marco Bianchi" in record.data["normalized_text"]
    assert "Nora Rossi" in record.data["normalized_text"]
    names = [s["display_name"] for s in record.data["structure"]["students_detected"]]
    assert names == ["Marco Bianchi", "Nora Rossi"]


def test_docx_roster_extracts_paragraph_and_table_text():
    from io import BytesIO

    import docx

    document = docx.Document()
    document.add_paragraph("Class List")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "First"
    table.rows[0].cells[1].text = "Last"
    for first, last in [("Marco", "Bianchi"), ("Nora", "Rossi")]:
        cells = table.add_row().cells
        cells[0].text = first
        cells[1].text = last
    buffer = BytesIO()
    document.save(buffer)
    source, _ = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    word = SourceRecord({
        **source.data,
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "original_ext": ".docx",
    })

    record = _run(extract_document(word, buffer.getvalue()))

    assert "Class List" in record.data["normalized_text"]
    names = [s["display_name"] for s in record.data["structure"]["students_detected"]]
    assert names == ["Marco Bianchi", "Nora Rossi"]


# --- Job runner ----------------------------------------------------------------


def _seed_vault(root: Path) -> str:
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    vault.put_source(source, content, root=root)
    return source.source_id


def test_job_runs_to_done_and_writes_extraction(tmp_path):
    source_id = _seed_vault(tmp_path)
    job = _run(jobs.run_extraction_job(source_id, root=tmp_path))
    assert job["status"] == "done", job
    assert "2 students detected" in job["progress"]["detail"]
    assert (tmp_path / "extracted" / f"{source_id}.json").exists()
    on_disk = jobs.job_status(job["job_id"], root=tmp_path)
    assert on_disk["status"] == "done"


def test_crashed_job_resumes_to_done_without_partials(tmp_path):
    source_id = _seed_vault(tmp_path)
    # Simulate the app dying mid-job: a job record stuck in "running".
    stuck = jobs._new_job(source_id)
    stuck["status"] = "running"
    jobs._write_job(stuck, tmp_path)
    assert not (tmp_path / "extracted" / f"{source_id}.json").exists()

    resumed = _run(jobs.resume_pending(root=tmp_path))
    assert len(resumed) == 1
    assert resumed[0]["status"] == "done"
    assert (tmp_path / "extracted" / f"{source_id}.json").exists()
    # exactly one extraction file, no temp leftovers
    files = list((tmp_path / "extracted").iterdir())
    assert [f.name for f in files] == [f"{source_id}.json"]


def test_job_failure_is_honest(tmp_path):
    job = _run(jobs.run_extraction_job("SRC-DOES-NOT-EXIST", root=tmp_path))
    assert job["status"] == "failed"
    assert job["error"]


# --- Support xlsx (FIX_SUPPORT_XLSX_LENS_PARSING 2026-08-18) -------------------


def _support_workbook_bytes():
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Classroom Accommodations"
    ws1.append(["Student Name", "Accommodation"])
    ws1.append(["Marco Bianchi", "Extended time on assessments"])
    ws1.append(["Nora Rossi", "Preferential seating"])
    ws2 = wb.create_sheet("External Support")
    ws2.append(["Student Name", "Provider", "Service"])
    ws2.append(["Marco Bianchi", "Dr. Smith", "Speech therapy"])
    ws3 = wb.create_sheet("Educational Evaluation")
    ws3.append(["Student Name", "Date", "Summary"])
    ws3.append(["Nora Rossi", "2026-03-15", "Cognitive assessment complete"])
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _support_source() -> SourceRecord:
    source, _ = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    return SourceRecord({
        **source.data,
        "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "original_ext": ".xlsx",
        "original_filename": "3V ES Student Support .xlsx",
    })


def test_support_xlsx_detects_students_not_sheet_names():
    """Multi-sheet support xlsx: real students detected, sheet names are NOT."""
    record = _run(extract_document(_support_source(), _support_workbook_bytes()))

    names = [s["display_name"] for s in record.data["structure"]["students_detected"]]
    assert "Marco Bianchi" in names
    assert "Nora Rossi" in names
    assert "Classroom Accommodations" not in names
    assert "External Support" not in names
    assert "Educational Evaluation" not in names
    assert record.data["structure"]["document_type"] == "student_support"
    # Spans carry field hints mapping into lens profile fields
    hints = {s.get("field_hint") for s in record.data["spans"] if s.get("field_hint")}
    assert "learning_and_cognition" in hints
    assert "communication_and_language" in hints


def test_support_xlsx_passes_grounding_gate():
    """field_hint is additive — every span is still an exact slice and every
    detected student is supported by cited spans (nothing dropped)."""
    record = _run(extract_document(_support_source(), _support_workbook_bytes()))
    report = verify_extraction(record.data)
    assert report.ok, report.errors
    assert not report.dropped


def test_support_xlsx_spans_feed_lens_fields():
    """The field_hint routes each row into the mapped lens field."""
    from src.lingua_viva.docpipe.lens import _fields_for_span

    record = _run(extract_document(_support_source(), _support_workbook_bytes()))
    marco_spans = [
        s for s in record.data["spans"] if "Marco Bianchi" in str(s.get("text"))
    ]
    assert marco_spans
    for span in marco_spans:
        fields = _fields_for_span(str(span.get("text") or ""), span=span)
        assert fields == [span["field_hint"]]


def test_plain_roster_xlsx_still_uses_generic_path():
    """A single-sheet roster ("Sheet") must not flip to the support path."""
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["First", "Last"])
    ws.append(["Marco", "Bianchi"])
    buf = BytesIO()
    wb.save(buf)
    source, _ = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    xlsx = SourceRecord({
        **source.data,
        "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "original_ext": ".xlsx",
    })
    record = _run(extract_document(xlsx, buf.getvalue()))
    assert record.data["structure"]["document_type"] != "student_support"
    assert not any(s.get("field_hint") for s in record.data["spans"])


# --- STEP 1: structure survives extraction (SPEC_LV_UNIFIED_REAL_DATA_FIX §STEP 1)


SYNTHETIC_CORPUS = Path(__file__).parent / "fixtures" / "docpipe" / "synthetic-corpus"


def _xlsx_source(filename: str) -> SourceRecord:
    source, _ = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    return SourceRecord({
        **source.data,
        "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "original_ext": ".xlsx",
        "original_filename": filename,
    })


def test_xlsx_rows_become_individual_spans_not_one_sheet_blob():
    """The 3V-shaped fixture yields one span per row (7: header + 6
    students), not 1 span for the whole sheet — the L9 failure."""
    content = (SYNTHETIC_CORPUS / "synthetic_support_3v.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_support_3v.xlsx"), content))

    spans = record.data["spans"]
    assert len(spans) == 7, f"expected 7 row spans, got {len(spans)}"
    # each of the 6 student rows is its OWN span
    student_spans = [s for s in spans if s["row_index"] >= 2]
    assert len(student_spans) == 6
    assert len({s["span_id"] for s in student_spans}) == 6


def test_xlsx_spans_carry_sheet_row_and_column_identity():
    content = (SYNTHETIC_CORPUS / "synthetic_support_3v.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_support_3v.xlsx"), content))

    header = record.data["spans"][0]
    assert header["sheet"] == "Sheet1"
    assert header["row_index"] == 1
    assert [c["column"] for c in header["cells"]][:2] == ["A", "B"]
    assert header["cells"][0]["text"] == "Student"
    assert header["cells"][1]["text"] == "class"
    first_student = record.data["spans"][1]
    assert first_student["cells"][0]["column"] == "A"
    assert first_student["cells"][1]["text"] in ("V", "A")


def test_xlsx_multi_sheet_spans_keep_sheet_names():
    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))

    sheets = {s["sheet"] for s in record.data["spans"]}
    assert sheets == {"Grade 3", "Grade 6"}
    # row 1 (class names) and row 2 (teachers) survive as their own spans
    grade3 = [s for s in record.data["spans"] if s["sheet"] == "Grade 3"]
    assert grade3[0]["row_index"] == 1
    assert grade3[1]["row_index"] == 2
    # empty columns don't produce cells; column letters are the real ones
    assert all(c["text"] for s in record.data["spans"] for c in s["cells"])


def test_xlsx_row_spans_pass_grounding_gate():
    """Row spans stay byte-exact slices — the grounding gate must hold."""
    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))

    report = verify_extraction(record.data, apply_drops=False)
    assert report.ok, report.errors


def test_support_path_untouched_by_row_spans():
    """Multi-sheet SUPPORT workbooks keep the sheet-aware support parse
    (field_hint spans), not the generic row-span path."""
    record = _run(extract_document(_support_source(), _support_workbook_bytes()))
    assert record.data["structure"]["document_type"] == "student_support"
    assert any(s.get("field_hint") for s in record.data["spans"])
    assert not any("row_index" in s for s in record.data["spans"])


def test_xlsx_extractions_survive_the_vault_schema_gate(tmp_path):
    """Class lock: additive span metadata (field_hint from the support path,
    sheet/row_index/cells from STEP 1) must pass the vault's schema
    validation — the 2026-08-18 support fix shipped spans the vault refused
    ('Could not read this document'), which this test would have caught."""
    from src.lingua_viva.docpipe import vault

    # STEP 1 row-span path (generic xlsx)
    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))
    assert any("row_index" in s for s in record.data["spans"])
    vault.put_extraction(record, root=tmp_path)

    # support path (field_hint spans)
    record = _run(extract_document(_support_source(), _support_workbook_bytes()))
    assert any(s.get("field_hint") for s in record.data["spans"])
    vault.put_extraction(record, root=tmp_path)


# --- STEP 2: detect from structure, not text shape (SPEC §STEP 2) --------------


def test_structured_zero_student_files_yield_zero_detections():
    """Gate: curriculum + calendar fixtures → 0 detections. Story titles and
    calendar labels stop being students because they are not in a Student
    column — positional evidence, not blocklist additions."""
    for filename in ("synthetic_curriculum.xlsx", "synthetic_calendar.xlsx"):
        content = (SYNTHETIC_CORPUS / filename).read_bytes()
        record = _run(extract_document(_xlsx_source(filename), content))
        assert record.data["structure"]["students_detected"] == [], filename


def test_student_column_detects_abbreviated_names():
    """Gate: 3V fixture → 6. Abbreviated names ('Marco B-R') are students
    because they sit in the Student column — no full-bigram requirement."""
    content = (SYNTHETIC_CORPUS / "synthetic_support_3v.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_support_3v.xlsx"), content))

    students = record.data["structure"]["students_detected"]
    assert len(students) == 6, [s["display_name"] for s in students]
    assert all(s["evidence"] == "student_column" for s in students)
    names = {s["display_name"] for s in students}
    assert "Marco B-R" in names
    # every detection cites its own row span
    for student in students:
        assert student["span_ids"], student["display_name"]


def test_staff_block_above_header_is_not_students():
    """K-5 shape: staff first names sit ABOVE the Student/Accommodations
    header row — position says they are not students."""
    content = (SYNTHETIC_CORPUS / "synthetic_support_k5.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_support_k5.xlsx"), content))

    names = {s["display_name"] for s in record.data["structure"]["students_detected"]}
    assert names == {
        "Marco Bianchi", "Nora Rossi", "Sara Conti",
        "Giulia Riva", "Pietro Serra", "Leo Fontana", "Camilla Gatti",
    }


def test_per_class_support_sheet_without_name_header_detects_abbreviations(tmp_path):
    """K-5 holdout shape: class-code sheet, no labelled name column, abbreviated
    row keys, and support categories in the other columns. Staff/header/summary
    rows must not become students."""
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "2V"
    ws.append(["2V", "2025-26 historical support"])
    ws.append(["Include Specialist", "support staff"])
    ws.append(["Sofia M.", "Classroom accommodations: extra time", "Medical needs: none"])
    ws.append(["Marco B-R", "Internal support weekly", "End-of-year notes: decoding support"])
    ws.append(["Summary", "Two students continue with support"])
    other = wb.create_sheet("Calendar")
    other.append(["Buddy Readers", "Assembly Notes"])
    buf = BytesIO()
    wb.save(buf)

    record = _run(extract_document(_xlsx_source("k5-support.xlsx"), buf.getvalue()))
    students = record.data["structure"]["students_detected"]

    assert [s["display_name"] for s in students] == ["Sofia M.", "Marco B-R"]
    assert {s["evidence"] for s in students} == {"per_class_sheet_support"}
    assert {s["class"] for s in students} == {"2V"}
    assert all(s["source_rows"].startswith("2V row ") for s in students)
    assert record.data["structure"]["document_type"] == "student_support"
    vault.put_extraction(record, root=tmp_path)


def test_first_last_column_pair_joins_names():
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["First", "Last", "Notes"])
    ws.append(["Marco", "Bianchi", "reading"])
    ws.append(["Nora", "Rossi", ""])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("pair.xlsx"), buf.getvalue()))
    names = [s["display_name"] for s in record.data["structure"]["students_detected"]]
    assert names == ["Marco Bianchi", "Nora Rossi"]


def test_adult_name_columns_are_never_student_columns():
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Teacher Name", "Room"])
    ws.append(["Ilaria Moretti", "12"])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("teachers.xlsx"), buf.getvalue()))
    assert record.data["structure"]["students_detected"] == []


def test_unstructured_documents_keep_the_bigram_fallback():
    """Markdown lesson plans have no columns — the bigram path survives as
    the fallback, tagged as its own (lower-confidence) evidence class."""
    source, content = _source("source_lesson_plan_marco_nora.json", "lesson_plan_marco_nora.md")
    record = _run(extract_document(source, content))
    students = record.data["structure"]["students_detected"]
    assert students, "bigram fallback must still detect in unstructured docs"
    assert all(s["evidence"] == "bigram_fallback" for s in students)


def test_repeated_header_rows_are_structure_not_students():
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Student", "class"])
    ws.append(["Marco B-R", "V"])
    ws.append(["Student", "class"])  # page-break header repeat
    ws.append(["Nora R-S", "A"])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("repeat.xlsx"), buf.getvalue()))
    names = [s["display_name"] for s in record.data["structure"]["students_detected"]]
    assert names == ["Marco B-R", "Nora R-S"]


def test_prose_mentioning_students_is_not_a_name_column():
    """Exact-label rule: a header IS a student-name label, never prose that
    contains one — mottos ('...per studenti...'), plan columns ('Student
    Support Plan'), and descriptions ('Gli studenti sono...') are not name
    columns. This is what produced the curriculum/3V false positives."""
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Niente senza gioia · per studenti curiosi", "Unità"])
    ws.append(["Le Stagioni", "Uno"])
    ws2 = wb.create_sheet("Plan")
    ws2.append(["Student", "Student Support Plan", "Gli studenti sono qui descritti"])
    ws2.append(["Marco B-R", "Piano Alfa", "Nota Lunga"])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("prose.xlsx"), buf.getvalue()))
    names = [s["display_name"] for s in record.data["structure"]["students_detected"]]
    # only the exact "Student" column detects; plan/prose columns never do
    assert names == ["Marco B-R"]


def test_nome_cognome_pair_and_lone_nome_column():
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Nome", "Cognome"])
    ws.append(["Marco", "Bianchi"])
    ws2 = wb.create_sheet("Lone")
    ws2.append(["Nome", "Classe"])
    ws2.append(["Nora Rossi", "3V"])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("nome.xlsx"), buf.getvalue()))
    names = {s["display_name"] for s in record.data["structure"]["students_detected"]}
    assert names == {"Marco Bianchi", "Nora Rossi"}


# --- STEP 4: class membership from the pair shape (SPEC §STEP 4, L4) -----------


def test_class_pair_sheets_detect_students_with_class_membership():
    """Gate: the class-list fixture (NO header labels — class name over each
    Last|First pair, teacher row under it) yields every roster student, each
    carrying class / grade / teacher_attribution. The teacher wants her ~39,
    not 400."""
    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))

    students = record.data["structure"]["students_detected"]
    assert len(students) == 19, [s["display_name"] for s in students]
    assert all(s["evidence"] == "student_column" for s in students)
    by_class: dict = {}
    for s in students:
        by_class.setdefault(s["class"], []).append(s)
    assert {c: len(v) for c, v in by_class.items()} == {
        "Verdi": 6, "Arancioni": 5, "Blu": 4, "Gialli": 4,
    }
    # membership metadata is complete and correct per pair
    verdi = by_class["Verdi"][0]
    assert verdi["grade"] == "Grade 3"
    assert verdi["teacher_attribution"] == "Ilaria Moretti"
    blu = by_class["Blu"][0]
    assert blu["grade"] == "Grade 6"
    assert blu["teacher_attribution"] == "Chiara De Luca"
    # Last|First pairs join as "First Last"
    names = {s["display_name"] for s in students}
    assert "Marco Bianchi" in names
    assert "Anna (Annie) Villa" in names  # nickname parentheses retained
    # every detection cites its row span
    assert all(s["span_ids"] for s in students)


def test_teacher_row_is_attribution_never_a_student():
    """Position says row 2 is the teacher row — teachers must never appear
    in the student set (the 'staff detected as students' failure class)."""
    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))

    names = {s["display_name"] for s in record.data["structure"]["students_detected"]}
    teachers = {"Ilaria Moretti", "Davide Colombo", "Chiara De Luca", "Stefano Ricci"}
    assert not names & teachers
    # ...but they ARE the attribution on their class's students
    attributions = {
        s["teacher_attribution"]
        for s in record.data["structure"]["students_detected"]
    }
    assert attributions == teachers


def test_title_rows_inside_roster_columns_are_not_students():
    """'1/2 groups...' title rows live INSIDE the roster columns but fill
    only one cell of the pair — excluded structurally, not by text matching
    or blocklist (§5)."""
    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))

    names = {s["display_name"] for s in record.data["structure"]["students_detected"]}
    assert not any("groups" in name for name in names)


# --- F1 (SPEC_LV_DEMO_EVE_FIX_2026-08-19): roster-block segmentation ----------


_ROSTER_NAMES = {
    # Verdi (Grade 3, cols A/B)
    "Marco Bianchi", "Nora Rossi", "Sara Conti",
    "Luca Ferri", "Elisa Greco", "Tommaso Marini",
    # Arancioni (Grade 3, cols C/D)
    "Giulia Riva", "Pietro Serra", "Anna (Annie) Villa",
    "Leo Fontana", "Camilla Gatti",
    # Blu (Grade 6, cols A/B)
    "Matteo Barbieri", "Sofia Caruso", "Emma Longo", "Diego Marchetti",
    # Gialli (Grade 6, cols C/D)
    "Alice Monti", "Filippo Neri", "Viola Pellegrini", "Oscar Sanna",
}


def test_group_table_blocks_never_create_students():
    """CLASS LOCK (F1, demo-blocking Finding 1): the fixture stacks two
    '1/2 groups' tables below each roster in the SAME column pairs, each
    column holding a COMPLETE name. The surname+firstname join applied to
    those blocks fused two real children into one non-existent student
    (18 → 35 on the real file). The join applies ONLY inside the roster
    block: detection must return exactly the roster names — zero fused."""
    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))

    names = {s["display_name"] for s in record.data["structure"]["students_detected"]}
    assert names == _ROSTER_NAMES
    # the specific fabrication shape: "First1 Last1 First2 Last2"
    fused = {n for n in names if len(n.split()) >= 4 and "(" not in n}
    assert fused == set(), fused


def test_roster_block_provenance_on_every_class_pair_student():
    """F1.4: preview must show WHICH rows the names came from, so a wrong
    segmentation is visible before confirm. Every class-pair student carries
    source_rows naming sheet, roster row range and column pair."""
    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))

    students = record.data["structure"]["students_detected"]
    assert students
    for s in students:
        assert "source_rows" in s, s["display_name"]
        assert s["grade"] in s["source_rows"]
        assert "rows 3–" in s["source_rows"]  # rosters start at row 3
        assert "cols" in s["source_rows"]


def test_asymmetric_rosters_segment_per_pair():
    """Two classes side by side with DIFFERENT roster lengths (real shape:
    17 + 19): the shorter pair's block ends where its columns stop filling;
    the longer pair keeps its full roster. Rows where only the other pair
    is filled never leak across."""
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Grade 4"
    ws.append(["Rossi Class", None, "Blu Class"])            # row 1: class names
    ws.append(["Teacher Uno", None, "Teacher Due"])          # row 2: teachers
    ws.append(["Alfa", "Anna", "Delta", "Dario"])            # rosters
    ws.append(["Beta", "Bruno", "Echo", "Enea"])
    ws.append([None, None, "Foxtrot", "Fabio"])              # only right pair
    ws.append([None, None, "Golf", "Gaia"])
    ws.append([])                                             # blank separator
    ws.append(["1/2 groups for Music/Homeroom", None, "1/2 groups for Music/Homeroom"])
    ws.append(["Anna Alfa", "Bruno Beta", "Dario Delta", "Enea Echo"])
    ws.append([None, None, "Fabio Foxtrot", "Gaia Golf"])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("asym.xlsx"), buf.getvalue()))

    students = record.data["structure"]["students_detected"]
    names = {s["display_name"] for s in students}
    assert names == {
        "Anna Alfa", "Bruno Beta",
        "Dario Delta", "Enea Echo", "Fabio Foxtrot", "Gaia Golf",
    }
    by_class: dict = {}
    for s in students:
        by_class.setdefault(s["class"], set()).add(s["display_name"])
    assert by_class["Rossi Class"] == {"Anna Alfa", "Bruno Beta"}
    assert by_class["Blu Class"] == {
        "Dario Delta", "Enea Echo", "Fabio Foxtrot", "Gaia Golf",
    }


def test_title_and_subtitle_over_a_grid_is_not_a_class_pair():
    """The exact real-corpus regression class: a sheet whose first two rows
    both fill only column A (title + subtitle) above a dense grid must yield
    ZERO students — the ≥2-pairs rule and the allowed-columns row guard fail
    closed. First cut of STEP 4 turned these into +217/+35 FP."""
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Mappatura"
    ws.append(["Mappatura UOI"])          # title — cols {A}
    ws.append(["Anno Scolastico"])        # subtitle — cols {A} == row 1
    ws.append(["Unità Uno", "Le Stagioni", "Autunno Inverno"])
    ws.append(["Unità Due", "La Famiglia", "Nonna Nonno"])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("grid.xlsx"), buf.getvalue()))
    assert record.data["structure"]["students_detected"] == []


def test_adjacent_heading_cells_are_a_grid_not_class_pairs():
    """A heading row with ADJACENT filled cells (A and B) is a table header,
    not class names over Last|First pairs — even if row 2 fills the same
    columns."""
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Giorno Uno", "Aula Musica", "Lunedì Mattina", "Orario Pieno"])
    ws.append(["Giorno Due", "Aula Arte", "Martedì Mattina", "Orario Ridotto"])
    ws.append(["Giorno Tre", "Palestra Grande", "Mercoledì Mattina", "Orario Pieno"])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("cycle.xlsx"), buf.getvalue()))
    assert record.data["structure"]["students_detected"] == []


def test_label_header_wins_over_pair_shape_and_class_column_scopes():
    """On a sheet WITH an explicit Student/class label header, the label
    rule wins (never both), and the class column attaches membership
    metadata without ever being a name source."""
    content = (SYNTHETIC_CORPUS / "synthetic_support_3v.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_support_3v.xlsx"), content))

    students = record.data["structure"]["students_detected"]
    assert len(students) == 6
    assert {s["class"] for s in students} == {"V", "A"}
    assert sum(1 for s in students if s["class"] == "V") == 3
    # class codes are metadata, never students
    names = {s["display_name"] for s in students}
    assert not names & {"V", "A"}


def test_lone_class_column_never_declares_a_header_row():
    """A row whose only recognized label is 'class'/'classe' is scoping
    metadata, not a student-name header — it must not turn every row below
    into students."""
    from io import BytesIO

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Materia", "classe", "Aula"])
    ws.append(["Lettura Guidata", "3V", "12"])
    ws.append(["Scrittura Creativa", "4A", "14"])
    buf = BytesIO()
    wb.save(buf)
    record = _run(extract_document(_xlsx_source("materie.xlsx"), buf.getvalue()))
    assert record.data["structure"]["students_detected"] == []


def test_class_pair_extractions_survive_the_vault_schema_gate(tmp_path):
    """The STEP 1 lesson, re-applied: new detection metadata (class / grade /
    teacher_attribution) must pass the vault's schema validation — additive
    keys that the schema refuses would break every real ingest."""
    from src.lingua_viva.docpipe import vault

    content = (SYNTHETIC_CORPUS / "synthetic_class_list.xlsx").read_bytes()
    record = _run(extract_document(_xlsx_source("synthetic_class_list.xlsx"), content))
    students = record.data["structure"]["students_detected"]
    assert all("class" in s and "grade" in s and "teacher_attribution" in s
               for s in students)
    vault.put_extraction(record, root=tmp_path)


# --- STEP 11 (C4): metadata from structure, not first-plausible-line ----------


def test_title_is_the_line_above_the_label_block_not_the_letterhead():
    """A labelled lesson export (synthetic mirror of the real IB PDF shape):
    letterhead + template header sit ABOVE the true title, which sits
    immediately above the metadata labels. First-plausible-line picked the
    letterhead; structure picks the title."""
    from src.lingua_viva.docpipe.extract import parse_lesson_metadata

    text = (
        "Sunrise International School School code: 012345\n"
        "Learning experience\n"
        "Instinct and Reason - How Animals React\n"
        "Unit: Diversity and emotions\n"
        "Author: Grade 3 Blue Teachers\n"
        "Subjects: Social Emotional Learning , Science , English\n\n"
        "Learning intentions\n"
        "1. transfer observations of animal reactions to stimuli."
    )
    meta = parse_lesson_metadata(text)
    assert meta["title"] == "Instinct and Reason - How Animals React"
    curriculum = meta["curriculum"]
    assert curriculum["grade"] == "G3"
    assert curriculum["unit"] == "Diversity and emotions"
    assert curriculum["subject"] == "Social Emotional Learning"


def test_grade_label_and_author_grade_patterns_normalize_to_grade_bands():
    from src.lingua_viva.docpipe.extract import _grade_from_text, parse_lesson_metadata

    assert _grade_from_text("3") == "G3"
    assert _grade_from_text("Grade 3 Blue Teachers") == "G3"
    assert _grade_from_text("MYP2 English") == "MYP2"
    assert _grade_from_text("pyp4") == "PYP4"
    assert _grade_from_text("Blue Teachers") is None

    meta = parse_lesson_metadata("My Lesson\nGrade: 4\nUnit: Water cycles")
    assert meta["curriculum"]["grade"] == "G4"
    # An Author line with no grade token adds nothing.
    meta = parse_lesson_metadata("My Lesson\nAuthor: Blue Teachers\nUnit: Water cycles")
    assert "grade" not in meta["curriculum"]


def test_title_detection_keeps_markdown_and_unlabelled_behavior():
    """Structure only kicks in when a label block exists: markdown headings
    still win, unlabelled documents keep the first-plausible-line heuristic,
    and Class-label grade parsing is unchanged."""
    from src.lingua_viva.docpipe.extract import parse_lesson_metadata

    assert parse_lesson_metadata("# Poetry Lesson\n\nRead the poem.")["title"] == "Poetry Lesson"
    assert parse_lesson_metadata("Weather Words\n\nA list of words.")["title"] == "Weather Words"
    meta = parse_lesson_metadata("Class: MYP2 English\nUnit: Migration stories")
    assert meta["curriculum"]["grade"] == "MYP2"
    assert meta["curriculum"]["subject"] == "English"


def test_implausible_line_above_labels_falls_back_not_letterhead_walk():
    """If the nearest line above the label block is not a plausible title
    (too long / ends like a sentence), detection must NOT keep walking up
    into the letterhead — it falls back to the old heuristic."""
    from src.lingua_viva.docpipe.extract import parse_lesson_metadata

    text = (
        "Sunrise Letterhead\n"
        "This line ends like a sentence and cannot be a title.\n"
        "Unit: Diversity and emotions"
    )
    meta = parse_lesson_metadata(text)
    assert meta["title"] == "Sunrise Letterhead"
