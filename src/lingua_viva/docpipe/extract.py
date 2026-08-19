"""Grounded document extraction (T3, SPEC_T3_EXTRACTION_2026-08-04).

Two layers:
- a deterministic core that always runs (offline is a supported state):
  normalization with stable character offsets, paragraph spans, structure,
  and student detection — everything is COPIED from the document, never
  generated;
- an optional local-model enrichment pass whose every claim is mechanically
  verified against its cited span by grounding_docs before inclusion.
  Unsupported claims are DROPPED, never demoted (T3 prompt rule).
"""
from __future__ import annotations

import json
import re
import unicodedata
from datetime import datetime, timezone
from typing import Any, Optional

from .contracts import ExtractionRecord, SourceRecord
from .grounding_docs import name_tokens, span_contains_any_token, verify_extraction
from .model import ModelClient

EXTRACTOR_NAME = "docpipe.extract"
EXTRACTOR_VERSION = "1.1"

TEXT_MIMES = {"text/markdown", "text/plain", "text/csv", "text/x-markdown"}
TEXT_EXTS = {".md", ".markdown", ".txt", ".csv"}
PDF_MIMES = {"application/pdf"}
SPREADSHEET_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
SPREADSHEET_EXTS = {".xlsx"}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
DOCX_EXTS = {".docx"}

MODEL_STUDENT_CONFIDENCE = 0.7
VERBATIM_STUDENT_CONFIDENCE = 0.99
MODEL_MAX_SPANS = 60

# Capitalized words that are never a person's name in teaching documents.
# Keeps the First-Last bigram detector from inventing students out of
# headings ("Italian Literature") or labels ("Teacher Note").
_NAME_BLOCKLIST = {
    "italian", "english", "french", "spanish", "german", "literature",
    "language", "student", "students", "work", "sample", "task", "date",
    "class", "unit", "teacher", "note", "notes", "learning", "goal",
    "main", "activity", "exit", "ticket", "warm", "up", "poetry",
    "personal", "identity", "voice", "grade", "school", "lesson", "plan",
    "report", "progress", "term", "roster", "list", "level", "reading",
    "writing", "speaking", "listening", "differentiation", "next", "step",
    "first", "last", "name", "surname", "given", "family",
    "monday", "tuesday", "wednesday", "thursday", "friday", "saturday",
    "sunday", "january", "february", "march", "april", "may", "june",
    "july", "august", "september", "october", "november", "december",
    # Education/support vocabulary — sheet names and section headers in
    # support documents ("Classroom Accommodations", "External Support",
    # "Tool Kit") are two capitalized words and would otherwise be
    # detected as students (FIX_SUPPORT_XLSX_LENS_PARSING 2026-08-18).
    "classroom", "accommodations", "accommodation", "external", "support",
    "toolkit", "kit", "tool", "evaluation", "educational", "counselor",
    "counselling", "counseling", "therapy", "therapist", "goals", "coast",
    "cognitive", "behavioral", "behavioural", "social", "emotional",
    "sensory", "physical", "motor", "attendance", "engagement",
    "communication", "executive", "functioning", "strategies", "strengths",
    "enrichment", "advanced", "intervention", "assessment", "diagnostic",
    "services", "referral", "placement", "program", "programme",
    "individual", "individualized", "education", "special", "needs",
    "overview", "summary", "general", "information", "details",
    "west", "east", "north", "south",
}

# Sheet-name → lens profile-field mapping for multi-sheet support xlsx
# files. Keys are matched against the lowercased, space-collapsed sheet
# title (exact first, then longest-key substring). Values are
# docpipe.lens PROFILE_FIELDS ids, verbatim.
_SHEET_FIELD_MAP = {
    "classroom accommodations": "learning_and_cognition",
    "accommodations": "learning_and_cognition",
    "educational evaluation": "learning_and_cognition",
    "evaluation": "learning_and_cognition",
    "assessment": "learning_and_cognition",
    "goals": "learning_and_cognition",
    "coast goals": "learning_and_cognition",
    "external support": "communication_and_language",
    "communication": "communication_and_language",
    "speech": "communication_and_language",
    "executive functioning": "executive_functioning",
    "executive": "executive_functioning",
    "social skills": "social_skills",
    "social": "social_skills",
    "emotional regulation": "emotional_regulation",
    "emotional": "emotional_regulation",
    "behavior": "emotional_regulation",
    "physical sensory": "physical_sensory_needs",
    "sensory": "physical_sensory_needs",
    "physical": "physical_sensory_needs",
    "attendance": "attendance_and_engagement",
    "engagement": "attendance_and_engagement",
    "toolkit": "strategies_trialed",
    "tool kit": "strategies_trialed",
    "strategies": "strategies_trialed",
    "strengths": "academic_strengths",
}
# At least this many sheets must map before we treat the workbook as a
# structured support document (a lone "Strengths" tab in a roster file
# must not flip the whole extraction path).
_SUPPORT_SHEET_MATCH_THRESHOLD = 2

_IT_STOPWORDS = {"il", "la", "che", "di", "e", "un", "una", "per", "con", "del", "della", "gli", "sono"}
_EN_STOPWORDS = {"the", "and", "of", "to", "a", "in", "that", "with", "for", "is", "are", "on"}


async def extract_document(
    source: SourceRecord,
    content: bytes,
    *,
    model_client: ModelClient | None = None,
) -> ExtractionRecord:
    warnings: list[str] = []
    # Multi-sheet support xlsx gets a structured, sheet-aware parse first
    # (per-student rows → field-hinted spans). Anything else — including
    # xlsx rosters whose sheet names don't match the support map — falls
    # through to the generic paragraph path.
    support = None
    row_extract = None
    mime = str(source.data.get("mime") or "").lower()
    ext = str(source.data.get("original_ext") or "").lower()
    if mime in SPREADSHEET_MIMES or ext in SPREADSHEET_EXTS:
        support = _xlsx_support_extract(source, content)
        if support is None:
            # STEP 1 (SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19): structure
            # survives extraction — one span per row, column identity kept.
            row_extract = _xlsx_row_spans(content)
    if support is not None:
        text = support["text"]
        spans = support["spans"]
        structure = support["structure"]
    else:
        if row_extract is not None:
            text, spans = row_extract
        else:
            text = _normalize(source, content)
            spans = _build_spans(text)
        structure = _build_structure(text, spans)
    model_used: Optional[str] = None

    if model_client is not None:
        model_used, model_warnings = await _model_enrich_students(
            model_client, spans, structure["students_detected"]
        )
        warnings.extend(model_warnings)
    else:
        warnings.append("model_enrichment_unavailable:no local model client")

    data: dict[str, Any] = {
        "schema_version": "docpipe.extraction.v1",
        "source_id": source.source_id,
        "source_sha256": str(source.data.get("sha256") or ""),
        "extracted_at": _now(),
        "extractor": {
            "name": EXTRACTOR_NAME,
            "version": EXTRACTOR_VERSION,
            "model": model_used,
        },
        "mime": str(source.data.get("mime") or ""),
        "language": _detect_language(text),
        "normalized_text": text,
        "spans": spans,
        "structure": structure,
        "warnings": warnings,
    }
    # The mechanical grounding gate runs INSIDE extraction too (jobs.py runs
    # it again before the vault write): anything unsupported is dropped here
    # so no caller can ever observe an unverified extraction.
    report = verify_extraction(data, apply_drops=True)
    if report.dropped:
        data["warnings"].extend(report.dropped)
    return ExtractionRecord(data)


# --- Normalization (stable offsets — computed once, never re-derived) --------


def _normalize(source: SourceRecord, content: bytes) -> str:
    mime = str(source.data.get("mime") or "").lower()
    ext = str(source.data.get("original_ext") or "").lower()
    if mime in TEXT_MIMES or ext in TEXT_EXTS or mime.startswith("text/"):
        raw = content.decode("utf-8", errors="replace")
    elif mime in PDF_MIMES or ext == ".pdf":
        raw = _pdf_text(content)
    elif mime in SPREADSHEET_MIMES or ext in SPREADSHEET_EXTS:
        raw = _xlsx_text(content)
    elif mime in DOCX_MIMES or ext in DOCX_EXTS:
        raw = _docx_text(content)
    else:
        raise ValueError(
            f"unsupported format for extraction: {mime or ext or 'unknown'} — "
            "supported today: markdown, plain text, csv, pdf, xlsx, docx"
        )
    return _normalize_text(raw)


def _normalize_text(raw: str) -> str:
    text = raw.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")
    lines = [line.rstrip() for line in text.split("\n")]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip("\n")
    return text + "\n\n" if text else ""


def _pdf_text(content: bytes) -> str:
    import io

    try:
        import pdfplumber
    except ImportError as error:
        raise ValueError(
            "PDF support is not available on this computer (pdfplumber missing)"
        ) from error
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    text = "\n\n".join(page for page in pages if page.strip())
    if not text.strip():
        raise ValueError("no extractable text found in this PDF")
    return text


def _xlsx_text(content: bytes) -> str:
    import io

    try:
        import openpyxl
    except ImportError as error:
        raise ValueError(
            "Excel support is not available on this computer (openpyxl missing)"
        ) from error
    workbook = openpyxl.load_workbook(
        io.BytesIO(content), read_only=True, data_only=True
    )
    sheets: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = [
                    str(cell).strip()
                    for cell in row
                    if cell is not None and str(cell).strip()
                ]
                if cells:
                    rows.append(" ".join(cells))
            if rows:
                sheets.append("\n".join(rows))
    finally:
        workbook.close()
    text = "\n\n".join(sheets)
    if not text.strip():
        raise ValueError("no extractable text found in this Excel file")
    return text


def _xlsx_row_spans(content: bytes) -> Optional[tuple[str, list[dict[str, Any]]]]:
    """STEP 1 (L9): structured extraction for generic xlsx — one span per
    non-empty ROW, with sheet name, row index, and per-cell column identity
    retained as additive span keys ({"sheet", "row_index", "cells"}).

    The span text stays an exact slice of the normalized text (built through
    _build_spans), so the grounding gate holds by construction — same
    additive-metadata pattern as _xlsx_support_extract's field_hint.

    Returns (text, spans) or None (unreadable workbook / no content /
    span↔row pairing broke) — the caller falls back to the flat-text path.

    Interpretation of WHICH row is a header / teacher row / student row is
    deliberately NOT done here: STEP 1 preserves structure, STEP 2/4 read it.
    """
    import io

    try:
        import openpyxl
        from openpyxl.utils import get_column_letter
    except ImportError:
        return None
    try:
        workbook = openpyxl.load_workbook(
            io.BytesIO(content), read_only=True, data_only=True
        )
    except Exception:
        return None

    blocks: list[str] = []
    metas: list[dict[str, Any]] = []
    try:
        for sheet in workbook.worksheets:
            for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                cells: list[dict[str, str]] = []
                for col_index, cell in enumerate(row, start=1):
                    value = (
                        re.sub(r"\s+", " ", str(cell).strip()) if cell is not None else ""
                    )
                    if value:
                        cells.append({"column": get_column_letter(col_index), "text": value})
                if not cells:
                    continue
                blocks.append(" ".join(cell["text"] for cell in cells))
                metas.append({
                    "sheet": sheet.title,
                    "row_index": row_index,
                    "cells": cells,
                })
    finally:
        workbook.close()

    if not blocks:
        return None
    text = _normalize_text("\n\n".join(blocks))
    spans = _build_spans(text)
    if len(spans) != len(blocks):
        # a row collapsed or split unexpectedly — the span↔row pairing would
        # be wrong; fall back rather than attach lying metadata
        return None
    for span, meta in zip(spans, metas):
        span.update(meta)
    return text, spans


def _sheet_field(sheet_title: str) -> Optional[str]:
    """Map a sheet title to a lens profile field, or None."""
    title = re.sub(r"\s+", " ", sheet_title.strip().lower())
    if title in _SHEET_FIELD_MAP:
        return _SHEET_FIELD_MAP[title]
    # Substring fallback ("Toolkit Is" → "toolkit"), longest key first so
    # "tool kit" beats "tool"-style partials.
    for key in sorted(_SHEET_FIELD_MAP, key=len, reverse=True):
        if key in title:
            return _SHEET_FIELD_MAP[key]
    return None


def _xlsx_support_extract(
    source: SourceRecord, content: bytes
) -> Optional[dict[str, Any]]:
    """Structured parse for multi-sheet student-support workbooks.

    Returns {"text", "spans", "structure"} when the workbook looks like a
    support document (>= _SUPPORT_SHEET_MATCH_THRESHOLD sheets map to lens
    fields), otherwise None — the caller falls back to the generic path.
    Every span is an exact slice of the normalized text (built through
    _build_spans), so the grounding gate holds by construction; the added
    "field_hint" key is additive and routes the entry into the right lens
    field in docpipe.lens.
    """
    import io

    try:
        import openpyxl
    except ImportError:
        return None
    try:
        workbook = openpyxl.load_workbook(
            io.BytesIO(content), read_only=True, data_only=True
        )
    except Exception:
        return None

    matched: list[tuple[str, str, list[tuple[Any, ...]]]] = []
    try:
        sheet_fields = [
            (sheet.title, _sheet_field(sheet.title)) for sheet in workbook.worksheets
        ]
        if sum(1 for _, field in sheet_fields if field) < _SUPPORT_SHEET_MATCH_THRESHOLD:
            return None
        for sheet in workbook.worksheets:
            field = _sheet_field(sheet.title)
            if not field:
                continue
            rows = [tuple(row) for row in sheet.iter_rows(values_only=True)]
            matched.append((sheet.title, field, rows))
    finally:
        workbook.close()

    def _clean(cell: Any) -> str:
        return re.sub(r"\s+", " ", str(cell).strip()) if cell is not None else ""

    blocks: list[str] = []
    field_hints: list[str] = []
    students_order: list[str] = []
    students: dict[str, dict[str, Any]] = {}
    sheet_blocks: dict[str, list[int]] = {}

    for sheet_title, field, rows in matched:
        non_empty = [row for row in rows if any(_clean(cell) for cell in row)]
        if len(non_empty) < 2:
            continue
        header = non_empty[0]
        name_col = 0
        for idx, cell in enumerate(header):
            label = _clean(cell).lower()
            if "name" in label or "student" in label:
                name_col = idx
                break
        for row in non_empty[1:]:
            name_raw = _clean(row[name_col]) if name_col < len(row) else ""
            match = _NAME_BIGRAM.search(name_raw)
            if not match:
                continue
            first, last = match.group(1), match.group(2)
            if first.lower() in _NAME_BLOCKLIST or last.lower() in _NAME_BLOCKLIST:
                continue
            display_name = f"{first} {last}"
            rest = "; ".join(
                _clean(cell)
                for idx, cell in enumerate(row)
                if idx != name_col and _clean(cell)
            )
            if not rest:
                continue
            blocks.append(f"{display_name} — {sheet_title}: {rest}")
            field_hints.append(field)
            sheet_blocks.setdefault(sheet_title, []).append(len(blocks) - 1)
            student_id = f"student-{_slug(display_name)}"
            if student_id not in students:
                students[student_id] = {
                    "student_id": student_id,
                    "display_name": display_name,
                    "confidence": VERBATIM_STUDENT_CONFIDENCE,
                    "evidence": "student_column",
                    "span_ids": [],
                }
                students_order.append(student_id)

    if not blocks:
        return None

    text = _normalize_text("\n\n".join(blocks))
    spans = _build_spans(text)
    if len(spans) != len(blocks):
        # A block contained an unexpected blank line — the span↔hint
        # pairing would be wrong. Fall back to the generic path.
        return None
    for span, hint in zip(spans, field_hints):
        span["field_hint"] = hint

    for student_id in students_order:
        student = students[student_id]
        for span in spans:
            if student["display_name"] in str(span["text"]):
                student["span_ids"].append(span["span_id"])

    sections = []
    span_by_block = {i: span["span_id"] for i, span in enumerate(spans)}
    for sheet_title, indices in sheet_blocks.items():
        sections.append({
            "section_id": _slug(sheet_title) or "support",
            "heading": sheet_title,
            "span_ids": [span_by_block[i] for i in indices],
        })

    filename = str(source.data.get("original_filename") or "").strip()
    title = re.sub(r"\.xlsx$", "", filename, flags=re.IGNORECASE).strip() or None
    return {
        "text": text,
        "spans": spans,
        "structure": {
            "title": title,
            "document_type": "student_support",
            "sections": sections,
            "students_detected": [students[sid] for sid in students_order],
        },
    }


def _docx_text(content: bytes) -> str:
    import io

    try:
        import docx
    except ImportError as error:
        raise ValueError(
            "Word document support is not available on this computer (python-docx missing)"
        ) from error
    document = docx.Document(io.BytesIO(content))
    blocks: list[str] = []
    blocks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" ".join(cells))
    text = "\n".join(blocks)
    if not text.strip():
        raise ValueError("no extractable text found in this Word document")
    return text


def extract_plain_text(content: bytes, ext: str) -> str:
    """Extension-driven text extraction for course-library lesson files.

    Same readers as _normalize but keyed on extension alone (course-library
    files carry no mime record). Lets lesson_materials read pdf/docx/xlsx
    lesson content that a plain utf-8 read would crash on.
    """
    ext = ext.lower()
    if ext in TEXT_EXTS:
        raw = content.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        raw = _pdf_text(content)
    elif ext in SPREADSHEET_EXTS:
        raw = _xlsx_text(content)
    elif ext in DOCX_EXTS:
        raw = _docx_text(content)
    else:
        raise ValueError(
            f"unsupported lesson file format: {ext or 'unknown'} — "
            "supported: markdown, plain text, csv, pdf, xlsx, docx"
        )
    return _normalize_text(raw)


def parse_lesson_metadata(text: str) -> dict[str, Any]:
    """Deterministic (no-model) lesson metadata from extracted text.

    Reuses the extraction span machinery so Class/Unit/Task labels parse
    exactly the way full docpipe extraction parses them. Used by Prepare to
    auto-fill Grade/Unit/Topic from an uploaded lesson file.
    """
    spans = _build_spans(text)
    students = _detect_students(spans)
    curriculum = _detect_curriculum(spans)
    return {
        "title": _detect_title(spans),
        "document_type": _detect_document_type(spans, students),
        "curriculum": curriculum,
    }


# --- Spans -------------------------------------------------------------------


def _build_spans(text: str) -> list[dict[str, Any]]:
    """Paragraph spans. text == normalized_text[start:end] holds by
    construction: spans are slices of the canonical string."""
    spans: list[dict[str, Any]] = []
    index = 0
    cursor = 0
    length = len(text)
    while cursor < length:
        while cursor < length and text[cursor] == "\n":
            cursor += 1
        if cursor >= length:
            break
        end = text.find("\n\n", cursor)
        if end == -1:
            end = length
        block = text[cursor:end]
        stripped_end = cursor + len(block.rstrip("\n"))
        if text[cursor:stripped_end].strip():
            index += 1
            spans.append({
                "span_id": f"SPN-{index:04d}",
                "char_start": cursor,
                "char_end": stripped_end,
                "text": text[cursor:stripped_end],
            })
        cursor = end
    return spans


# --- Structure ---------------------------------------------------------------


def _build_structure(text: str, spans: list[dict[str, Any]]) -> dict[str, Any]:
    title = _detect_title(spans)
    # STEP 2: structured documents (row spans with column identity) use
    # positional evidence ONLY. The bigram regex survives strictly as a
    # fallback for unstructured documents (its output is a lower-confidence
    # evidence class, "bigram_fallback").
    if any("cells" in span for span in spans):
        students = _detect_students_structural(spans)
    else:
        students = _detect_students(spans)
    document_type = _detect_document_type(spans, students)
    sections = _build_sections(spans, document_type, title)
    structure: dict[str, Any] = {
        "title": title,
        "document_type": document_type,
        "sections": sections,
        "students_detected": students,
    }
    curriculum = _detect_curriculum(spans)
    if curriculum:
        structure["curriculum"] = curriculum
    return structure


def _detect_title(spans: list[dict[str, Any]]) -> Optional[str]:
    if not spans:
        return None
    first = str(spans[0]["text"]).split("\n")[0]
    if first.startswith("#"):
        return first.lstrip("#").strip()
    if len(first) <= 100 and not first.endswith((".", ":")):
        return first.strip()
    return None


def _labeled_value(span_text: str, label: str) -> Optional[str]:
    for line in span_text.split("\n"):
        if line.lower().startswith(label.lower() + ":"):
            return line.split(":", 1)[1].strip()
    return None


def _detect_document_type(spans: list[dict[str, Any]], students: list[dict[str, Any]]) -> str:
    joined = " ".join(str(s["text"]).lower() for s in spans)
    if ("roster" in joined or "class list" in joined) and len(students) >= 3:
        return "roster"
    if any(k in joined for k in ("learning goal", "warm-up", "warm up", "main activity", "differentiation")):
        return "lesson_plan"
    if "work sample" in joined or ("task:" in joined and students):
        return "student_work_sample"
    if any(k in joined for k in ("criterion", "descriptor", "rubric", "band ")):
        return "rubric"
    if "progress report" in joined or "term report" in joined:
        return "report"
    if any(k in joined for k in (
        "accommodation", "support plan", "student support",
        "educational evaluation", "external support",
    )):
        return "student_support"
    return "unknown"


def _build_sections(
    spans: list[dict[str, Any]],
    document_type: str,
    title: Optional[str],
) -> list[dict[str, Any]]:
    if not spans:
        return []
    title_span_ids = {spans[0]["span_id"]} if title else set()
    meta_ids: list[str] = []
    differentiation_ids: list[str] = []
    task_ids: list[str] = []
    teacher_note_ids: list[str] = []
    flow_ids: list[str] = []
    prose_ids: list[str] = []
    for span in spans:
        sid = span["span_id"]
        if sid in title_span_ids:
            continue
        text = str(span["text"])
        lower = text.lower()
        if _labeled_value(text, "Class") is not None:
            meta_ids.append(sid)
        elif lower.startswith("differentiation"):
            differentiation_ids.append(sid)
        elif _labeled_value(text, "Task") is not None:
            task_ids.append(sid)
        elif lower.startswith("teacher note"):
            teacher_note_ids.append(sid)
        elif document_type == "lesson_plan":
            flow_ids.append(sid)
        else:
            prose_ids.append(sid)
    sections: list[dict[str, Any]] = []
    if meta_ids:
        sections.append({"section_id": "lesson-meta", "heading": "Class, date, unit", "span_ids": meta_ids})
    if task_ids:
        sections.append({"section_id": "task", "heading": "Task", "span_ids": task_ids})
    if flow_ids:
        sections.append({"section_id": "lesson-flow", "heading": "Lesson flow", "span_ids": flow_ids})
    if prose_ids:
        sections.append({"section_id": "student-response", "heading": "Student response", "span_ids": prose_ids})
    if differentiation_ids:
        sections.append({"section_id": "differentiation", "heading": "Differentiation notes", "span_ids": differentiation_ids})
    if teacher_note_ids:
        sections.append({"section_id": "teacher-note", "heading": "Teacher note", "span_ids": teacher_note_ids})
    return sections


def _detect_curriculum(spans: list[dict[str, Any]]) -> dict[str, Any]:
    curriculum: dict[str, Any] = {}
    for span in spans:
        text = str(span["text"])
        class_value = _labeled_value(text, "Class")
        if class_value and "grade" not in curriculum and "subject" not in curriculum:
            parts = class_value.split(None, 1)
            if parts and re.fullmatch(r"(MYP\d{1,2}|PYP\d{1,2}|DP\d?|G\d{1,2})", parts[0]):
                curriculum["grade"] = parts[0]
                if len(parts) > 1:
                    curriculum["subject"] = parts[1]
            else:
                curriculum["subject"] = class_value
        unit_value = _labeled_value(text, "Unit")
        if unit_value and "unit" not in curriculum:
            curriculum["unit"] = unit_value
        task_value = _labeled_value(text, "Task")
        if task_value and "task" not in curriculum:
            curriculum["task"] = task_value
    return curriculum


def _detect_language(text: str) -> str:
    words = re.findall(r"[a-zà-ù]+", text.lower())
    it_hits = sum(1 for w in words if w in _IT_STOPWORDS)
    en_hits = sum(1 for w in words if w in _EN_STOPWORDS)
    return "it" if it_hits > en_hits else "en"


# --- Student detection (deterministic) ---------------------------------------


def _slug(name: str) -> str:
    folded = "".join(
        ch for ch in unicodedata.normalize("NFKD", name) if not unicodedata.combining(ch)
    )
    return re.sub(r"[^a-z0-9]+", "-", folded.lower()).strip("-")


_NAME_BIGRAM = re.compile(r"\b([A-ZÀ-Þ][a-zà-ÿ]+)\s+([A-ZÀ-Þ][a-zà-ÿ]+)\b")

# STEP 2 (SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19): detect from structure.
# A student is a value in a column whose header IS a student-name label —
# exact label match, never substring: prose that merely mentions students
# ("Gli studenti sono...", "Student Support Plan") is not a name column.
# "Say where to go, not where not to go."
_FULL_NAME_LABELS = {
    "student", "students", "student name", "students name", "student names",
    "name", "full name", "nome e cognome", "cognome e nome",
    "alunno", "alunna", "alunni", "alunne",
    "studente", "studenti", "studentessa", "studentesse",
}
_LAST_NAME_LABELS = {"last", "last name", "surname", "family name", "cognome"}
_FIRST_NAME_LABELS = {"first", "first name", "given name"}
# STEP 4 (L4): a class column is METADATA, never a name source — it scopes
# an import ("only my class"), it does not detect students.
_CLASS_LABELS = {"class", "classe"}
# "nome" is Italian for both "name" and "first name": it is a first-name
# column when a cognome column shares the header row, otherwise full-name
# (resolved in _sheet_student_columns).


def _column_role(label: str) -> Optional[str]:
    """Classify a header cell: "full" | "first" | "last" | "nome" | "class",
    or None."""
    label = re.sub(r"\s+", " ", str(label or "").strip().lower()).rstrip(":").strip()
    if not label:
        return None
    if label == "nome":
        return "nome"
    if label in _FULL_NAME_LABELS:
        return "full"
    if label in _LAST_NAME_LABELS:
        return "last"
    if label in _FIRST_NAME_LABELS:
        return "first"
    if label in _CLASS_LABELS:
        return "class"
    return None


def _column_index(letter: str) -> int:
    index = 0
    for char in letter:
        index = index * 26 + (ord(char) - ord("A") + 1)
    return index


def _column_letter(index: int) -> str:
    letters = ""
    while index > 0:
        index, remainder = divmod(index - 1, 26)
        letters = chr(ord("A") + remainder) + letters
    return letters


def _sheet_student_columns(
    spans: list[dict[str, Any]],
) -> dict[str, tuple[int, dict[str, str]]]:
    """Per sheet: (header_row_index, {column_letter: role}) from the first
    row whose cells carry a student-name concept."""
    columns: dict[str, tuple[int, dict[str, str]]] = {}
    for span in spans:
        sheet = span.get("sheet")
        if sheet is None or sheet in columns or "cells" not in span:
            continue
        roles = {}
        for cell in span["cells"]:
            role = _column_role(cell["text"])
            if role:
                roles[cell["column"]] = role
        # A header row must carry a NAME concept — a lone "class" column is
        # scoping metadata, not a student-name header (STEP 4).
        if any(role != "class" for role in roles.values()):
            # "Nome" is a first name only when paired with a Cognome column
            has_last = any(role == "last" for role in roles.values())
            roles = {
                column: ("first" if has_last else "full") if role == "nome" else role
                for column, role in roles.items()
            }
            columns[sheet] = (int(span["row_index"]), roles)
    return columns


def _sheet_class_pairs(spans: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    """STEP 4 (L4, SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19): grade-sheet
    class lists carry NO header labels — the SHAPE is the evidence. Row 1
    names the class above each Last|First column pair, row 2 names the
    teacher under it, and roster rows fill BOTH columns of a pair.

    Recognized per sheet only when:
    - the first two non-empty rows fill exactly the same columns,
    - no pair-start's right neighbour is filled in the heading row (a table
      with adjacent header cells — curriculum grids, cycle calendars — is
      NOT this shape), and
    - there are AT LEAST TWO pairs: a single stacked header pair is
      structurally indistinguishable from a title + subtitle above a
      two-column table, so it fails closed to zero students.
    Position says row 2 is a teacher row: teachers are attribution, never
    students.
    """
    by_sheet: dict[str, list[dict[str, Any]]] = {}
    for span in spans:
        sheet = span.get("sheet")
        if sheet is None or "cells" not in span:
            continue
        by_sheet.setdefault(sheet, []).append(span)
    pairs_by_sheet: dict[str, dict[str, Any]] = {}
    for sheet, sheet_spans in by_sheet.items():
        sheet_spans = sorted(sheet_spans, key=lambda span: int(span["row_index"]))
        if len(sheet_spans) < 3:
            continue
        heading, teacher_row = sheet_spans[0], sheet_spans[1]
        heading_cols = {cell["column"]: cell["text"] for cell in heading["cells"]}
        teacher_cols = {cell["column"]: cell["text"] for cell in teacher_row["cells"]}
        if not heading_cols or set(heading_cols) != set(teacher_cols):
            continue
        pairs: dict[str, dict[str, str]] = {}
        for column in sorted(heading_cols, key=_column_index):
            neighbour = _column_letter(_column_index(column) + 1)
            if neighbour in heading_cols:
                pairs = {}
                break  # adjacent header cells — a grid, not class pairs
            pairs[column] = {
                "neighbour": neighbour,
                "class": heading_cols[column],
                "teacher": teacher_cols[column],
            }
        if len(pairs) >= 2:
            pairs_by_sheet[sheet] = {
                "teacher_row": int(teacher_row["row_index"]),
                "pairs": pairs,
                # a roster row lives INSIDE the pair columns — any row that
                # fills a column outside them is table data, not a student
                "allowed_columns": {
                    column
                    for pair_start, pair in pairs.items()
                    for column in (pair_start, pair["neighbour"])
                },
            }
    return pairs_by_sheet


def _detect_students_structural(spans: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Positional student detection for structured (row-span) documents.

    No bigram requirement — abbreviated names ("Marco B-R") are students
    because they sit in the Student column. A structured document with no
    student-name column yields ZERO students: story titles and calendar
    labels stop being students because they are not in a Student column,
    not because of blocklist entries. Unstructured documents fall back to
    _detect_students (the bigram path) in _build_structure.
    """
    header_by_sheet = _sheet_student_columns(spans)
    class_pairs_by_sheet = {
        sheet: info
        for sheet, info in _sheet_class_pairs(spans).items()
        if sheet not in header_by_sheet  # an explicit label header wins
    }
    found: dict[str, dict[str, Any]] = {}
    order: list[str] = []

    def _add(display_name: str, span_id: str, meta: Optional[dict[str, str]] = None) -> None:
        display_name = re.sub(r"\s+", " ", display_name).strip()
        if not display_name:
            return
        # a repeated header row is structure, not a student
        if _column_role(display_name):
            return
        student_id = f"student-{_slug(display_name)}"
        if student_id not in found:
            record = {
                "student_id": student_id,
                "display_name": display_name,
                "confidence": VERBATIM_STUDENT_CONFIDENCE,
                "evidence": "student_column",
                "span_ids": [],
            }
            for key, value in (meta or {}).items():
                value = re.sub(r"\s+", " ", str(value or "")).strip()
                if value:
                    record[key] = value
            found[student_id] = record
            order.append(student_id)
        if span_id not in found[student_id]["span_ids"]:
            found[student_id]["span_ids"].append(span_id)

    for span in spans:
        sheet = span.get("sheet")
        if "cells" not in span:
            continue
        row_index = int(span["row_index"])
        if sheet in header_by_sheet:
            header_row, roles = header_by_sheet[sheet]
            if row_index <= header_row:
                continue
            first_part = ""
            last_part = ""
            class_value = ""
            full_names: list[str] = []
            for cell in span["cells"]:
                role = roles.get(cell["column"])
                if role == "full":
                    full_names.append(cell["text"])
                elif role == "first":
                    first_part = cell["text"]
                elif role == "last":
                    last_part = cell["text"]
                elif role == "class":
                    class_value = cell["text"]
            meta = {"class": class_value} if class_value else None
            for name in full_names:
                _add(name, span["span_id"], meta)
            if first_part or last_part:
                _add(f"{first_part} {last_part}".strip(), span["span_id"], meta)
        elif sheet in class_pairs_by_sheet:
            info = class_pairs_by_sheet[sheet]
            if row_index <= info["teacher_row"]:
                continue
            cells = {cell["column"]: cell["text"] for cell in span["cells"]}
            if any(column not in info["allowed_columns"] for column in cells):
                continue  # fills beyond the pairs — table data, not a roster row
            for column, pair in info["pairs"].items():
                last_part = str(cells.get(column) or "").strip()
                first_part = str(cells.get(pair["neighbour"]) or "").strip()
                # A roster row fills BOTH columns of the pair; a row with one
                # cell (scheduling notes inside the roster columns) is not a
                # student — structurally, no text matching.
                if last_part and first_part:
                    _add(f"{first_part} {last_part}", span["span_id"], {
                        "class": pair["class"],
                        "grade": str(sheet),
                        "teacher_attribution": pair["teacher"],
                    })
    return [found[student_id] for student_id in order]


def _detect_students(spans: list[dict[str, Any]]) -> list[dict[str, Any]]:
    found: dict[str, dict[str, Any]] = {}
    order: list[str] = []
    for span in spans:
        text = str(span["text"])
        for match in _NAME_BIGRAM.finditer(text):
            first, last = match.group(1), match.group(2)
            if first.lower() in _NAME_BLOCKLIST or last.lower() in _NAME_BLOCKLIST:
                continue
            display_name = f"{first} {last}"
            student_id = f"student-{_slug(display_name)}"
            if student_id not in found:
                found[student_id] = {
                    "student_id": student_id,
                    "display_name": display_name,
                    "confidence": VERBATIM_STUDENT_CONFIDENCE,
                    "evidence": "bigram_fallback",
                    "span_ids": [],
                }
                order.append(student_id)
    # Span membership: full name OR first name as a whole word.
    for student_id in order:
        student = found[student_id]
        first_name = student["display_name"].split()[0]
        for span in spans:
            text = str(span["text"])
            if student["display_name"] in text or re.search(
                rf"\b{re.escape(first_name)}\b", text
            ):
                student["span_ids"].append(span["span_id"])
    return [found[student_id] for student_id in order]


# --- Model enrichment (optional, verified, drop-on-failure) ------------------


_MODEL_SYSTEM_PROMPT = (
    "You identify student names in a teaching document. Reply with STRICT "
    'JSON only, exactly: {"students": [{"display_name": "First Last", '
    '"span_id": "SPN-0001"}], "not_students": [{"display_name": "First Last", '
    '"span_id": "SPN-0001", "reason": "short reason"}]} . In "students", only '
    "include real student names that appear verbatim inside the cited span. "
    'In "not_students", only include names from the detected-students list '
    "that the cited span shows are NOT students (a teacher, a school name, a "
    'note). If a list is empty, use []. No prose, no markdown fences.'
)


async def _model_enrich_students(
    model_client: ModelClient,
    spans: list[dict[str, Any]],
    students: list[dict[str, Any]],
) -> tuple[Optional[str], list[str]]:
    warnings: list[str] = []
    prompt_spans = spans[:MODEL_MAX_SPANS]
    if len(spans) > MODEL_MAX_SPANS:
        warnings.append(f"model_enrichment_truncated:{len(spans) - MODEL_MAX_SPANS} spans not sent")
    prompt = "Spans:\n" + "\n".join(
        f"[{span['span_id']}] {str(span['text'])[:500]}" for span in prompt_spans
    )
    # STEP 6 (SPEC_LV_UNIFIED_REAL_DATA_FIX_2026-08-19, L7): the model can
    # only dispute what it can see — the current detections ride along so
    # "not_students" claims have something concrete to point at.
    if students:
        prompt += "\n\nDetected students: " + ", ".join(
            str(student.get("display_name") or "") for student in students
        )
    model_used: Optional[str] = None
    parsed: Optional[dict[str, Any]] = None
    for attempt in range(2):
        try:
            result = await model_client.complete(
                prompt if attempt == 0 else prompt + "\n\nReply with the JSON object ONLY.",
                system_prompt=_MODEL_SYSTEM_PROMPT,
                max_tokens=500,
            )
        except Exception as error:
            warnings.append(f"model_enrichment_unavailable:{str(error)[:120]}")
            return None, warnings
        model_used = result.model_used or model_used
        # STEP 7 (L6): "none*" model_used means no model ever ran (privacy
        # refusal, no model installed). Reporting the true reason here is the
        # fix for the audit's misreport — this used to fall through to JSON
        # parsing and surface as "invalid JSON after retry".
        if result.error or str(result.model_used or "").startswith("none"):
            reason = result.error or f"no usable local model ({result.model_used})"
            warnings.append(f"model_enrichment_unavailable:{reason}")
            return (model_used if model_used not in (None, "none") else None), warnings
        parsed = _parse_model_json(result.content)
        if parsed is not None:
            break
    if parsed is None:
        warnings.append("model_enrichment_discarded:invalid JSON after retry")
        return model_used, warnings

    known_ids = {student["student_id"] for student in students}
    span_by_id = {span["span_id"]: span for span in spans}
    for claim in parsed.get("students", []):
        if not isinstance(claim, dict):
            continue
        display_name = str(claim.get("display_name") or "").strip()
        span_id = str(claim.get("span_id") or "").strip()
        span = span_by_id.get(span_id)
        if not display_name or span is None:
            warnings.append(
                f"grounding_dropped:model_student:{display_name or '?'}:span {span_id or '?'} does not exist"
            )
            continue
        tokens = name_tokens(display_name)
        # The mechanical support rule: every name token must appear in the
        # cited span. Fails → DROPPED, never demoted.
        if not tokens or not all(
            span_contains_any_token(str(span["text"]), [token]) for token in tokens
        ):
            warnings.append(
                f"grounding_dropped:model_student:{display_name}:cited span {span_id} does not contain the name"
            )
            continue
        student_id = f"student-{_slug(display_name)}"
        if student_id in known_ids:
            continue
        known_ids.add(student_id)
        first_name = display_name.split()[0]
        span_ids = [
            candidate["span_id"]
            for candidate in spans
            if display_name in str(candidate["text"])
            or re.search(rf"\b{re.escape(first_name)}\b", str(candidate["text"]))
        ]
        students.append({
            "student_id": student_id,
            "display_name": display_name,
            "confidence": MODEL_STUDENT_CONFIDENCE,
            "evidence": "model",
            "span_ids": span_ids or [span_id],
        })

    # STEP 6 (SPEC §STEP 6, L7): the veto channel. Additive-only enrichment
    # cannot converge on truth — the model may now DISPUTE a detection, but
    # a veto never deletes: it sets removal_proposed and the teacher rules
    # (review-gated downstream, exactly like additions are grounding-gated
    # here). The mechanical rule is symmetric with additions: the claim must
    # name a current detection, cite a real span, and that span must contain
    # the name — an ungrounded veto is DROPPED, never trusted.
    from src.lingua_viva.docpipe.identity import normalize_name

    detected_by_norm = {
        normalize_name(str(student.get("display_name") or "")): student
        for student in students
    }
    for claim in parsed.get("not_students", []) if isinstance(parsed.get("not_students"), list) else []:
        if not isinstance(claim, dict):
            continue
        display_name = str(claim.get("display_name") or "").strip()
        span_id = str(claim.get("span_id") or "").strip()
        reason = str(claim.get("reason") or "").strip()
        student = detected_by_norm.get(normalize_name(display_name))
        span = span_by_id.get(span_id)
        if student is None or span is None or not reason:
            warnings.append(
                f"grounding_dropped:model_veto:{display_name or '?'}:"
                "claim must name a current detection, cite a real span, and give a reason"
            )
            continue
        tokens = name_tokens(display_name)
        if not tokens or not all(
            span_contains_any_token(str(span["text"]), [token]) for token in tokens
        ):
            warnings.append(
                f"grounding_dropped:model_veto:{display_name}:cited span {span_id} does not contain the name"
            )
            continue
        student["removal_proposed"] = reason[:160]
        warnings.append(f"model_enrichment_veto:{display_name}:{reason[:120]}")
    return model_used, warnings


def _parse_model_json(content: str) -> Optional[dict[str, Any]]:
    """Hardened JSON parse (observe/classify lessons): strip fences, take the
    first balanced object only, discard on invalid."""
    text = str(content or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    start = text.find("{")
    if start == -1:
        return None
    depth = 0
    for index in range(start, len(text)):
        if text[index] == "{":
            depth += 1
        elif text[index] == "}":
            depth -= 1
            if depth == 0:
                try:
                    candidate = json.loads(text[start:index + 1])
                except json.JSONDecodeError:
                    return None
                return candidate if isinstance(candidate, dict) else None
    return None


def _now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
